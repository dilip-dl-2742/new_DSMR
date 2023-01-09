import json

def lambda_handler(event, context):    
    import boto3 
    import pytz
    from datetime import datetime, timedelta, timezone
    import dateutil.parser
    from dateutil.tz import tzutc
    import pandas as pd
    import matplotlib.pyplot as plt 
    from docx import Document
    from docx.shared import Inches
    import numpy as np
    from PIL import Image
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import smtplib 
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders
    
    
    est = pytz.timezone('US/Eastern')
    utc = pytz.utc
    fmt = '%Y-%m-%d %H:%M %Z%z'
    
    #Code for Input time for last 12 Hours. 
    
    todayUTC = datetime.now(timezone.utc).date()
    yesterdayUTC = todayUTC - timedelta(1)
    
    todayUTC = str(todayUTC)
    yesterdayUTC = str(yesterdayUTC)
    
    t_split = todayUTC.split('-')
    t_input_year = t_split[0]
    t_input_month = t_split[1]
    t_input_day = t_split[2]
    
    y_split = yesterdayUTC.split('-')
    y_input_year = y_split[0]
    y_input_month = y_split[1]
    y_input_day = y_split[2]
    
    
    if int(t_input_day) <= 9:
        t_input_day = t_input_day[1]
    else:
        pass
    if int(t_input_month) <= 9:
        t_input_month = t_input_month[1]
    else:
        pass 
    
    if int(y_input_day) <= 9:
        y_input_day = y_input_day[1]
    else:
        pass
    if int(y_input_month) <= 9:
        y_input_month = y_input_month[1]
    else:
        pass 
    
    t_input_day = int(t_input_day)
    t_input_month = int(t_input_month)
    t_input_year = int(t_input_year)
    
    y_input_day = int(y_input_day)
    y_input_month = int(y_input_month)
    y_input_year = int(y_input_year)
    
    #Code to get Metric Statistics.
    
    #------------------- RDS Writer Instance CPU Utilization ---------------------
    
    ac_id = ""
    as_key = ""
    #a_s_t = ""
    
    client = boto3.client('cloudwatch',aws_access_key_id = ac_id,aws_secret_access_key = as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    rds_1_1_response = client.get_metric_statistics(
        Namespace='AWS/RDS',
        MetricName='CPUUtilization',
        Dimensions=[
            {
                'Name': 'DBClusterIdentifier',
                'Value': 'prod-amc-db-cluster'
            },
            {
                'Name':'Role',
                'Value':'WRITER'
            },
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),#22,0
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),#10,30
        Period=3600,
        Statistics=['Average']
    )
    #print(rds_1_1_response)
    rds_1_1_a = (rds_1_1_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    rds_1_1_date_list = []
    rds_1_1_cpu_list = []
    for i in rds_1_1_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        rds_1_1_date_list.append(est_dt)
        rds_1_1_cpu_list.append(c)
        
    rds_1_1_d = {"Dates : " : rds_1_1_date_list, "CPU : " : rds_1_1_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    rds_1_1_df = pd.DataFrame(rds_1_1_d, columns=['Dates : ', 'CPU : '])
    rds_1_1_df.sort_values(by=['Dates : ','CPU : '], inplace = True)
    rds_1_1_Timestamp = rds_1_1_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    rds_1_1_Timestamp_list = []
    
    for j in rds_1_1_Timestamp:
        time = j.split()
        time = time[1]
        rds_1_1_Timestamp_list.append(time)
    rds_1_1_CPU_Utilization = rds_1_1_df['CPU : '].tolist()
    
    #-------------LINEGRAPH----------------------------------
    
    
    fig1 = plt.rcParams['font.size'] = '7'
    
    plt.plot(rds_1_1_Timestamp_list, rds_1_1_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('CPU Utilization of RDS Aurora PostgresSQL(Writer)', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('CPU_Utilization', fontsize = 10)
    plt.grid()
    plt.savefig('rds_1_1_CPU.png', dpi=100)
    plt.close()
    
    
    #------------------- RDS Writer Instance DB Connections -----------------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    rds_1_2_response = client.get_metric_statistics(
        Namespace='AWS/RDS',
        MetricName='DatabaseConnections',
        Dimensions=[
            {
                'Name': 'DBClusterIdentifier',
                'Value': 'prod-amc-db-cluster'
            },
            {
                'Name':'Role',
                'Value':'WRITER'
            },
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),#22,0
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    
    rds_1_2_a = (rds_1_2_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    rds_1_2_date_list = []
    rds_1_2_cpu_list = []
    for i in rds_1_2_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        rds_1_2_date_list.append(est_dt)
        rds_1_2_cpu_list.append(c)
        
    rds_1_2_d = {"Dates : " : rds_1_2_date_list, "DB Connections : " : rds_1_2_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    rds_1_2_df = pd.DataFrame(rds_1_2_d, columns=['Dates : ', 'DB Connections : '])
    rds_1_2_df.sort_values(by=['Dates : ','DB Connections : '], inplace = True)
    rds_1_2_Timestamp = rds_1_2_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    rds_1_2_Timestamp_list = []
    
    for j in rds_1_2_Timestamp:
        time = j.split()
        time = time[1]
        rds_1_2_Timestamp_list.append(time)
    rds_1_2_CPU_Utilization = rds_1_2_df['DB Connections : '].tolist()
    
    fig2 = plt.rcParams['font.size'] = '7'
    
    plt.plot(rds_1_2_Timestamp_list, rds_1_2_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('DB Connections for RDS Aurora PostgresSQL(Writer)', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('DB Connections', fontsize = 10)
    plt.grid()
    plt.savefig('rds_1_2_DB_Connection.png', dpi=100)
    plt.close()
    
    #-------------------------RDS Writer Insatance Freeable Memory---------------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    rds_1_3_response = client.get_metric_statistics(
        Namespace='AWS/RDS',
        MetricName='FreeableMemory',
        Dimensions=[
            {
                'Name': 'DBClusterIdentifier',
                'Value': 'prod-amc-db-cluster'
            },
            {
                'Name':'Role',
                'Value':'WRITER'
            },
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    
    rds_1_3_a = (rds_1_3_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    rds_1_3_date_list = []
    rds_1_3_cpu_list = []
    for i in rds_1_3_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        rds_1_3_date_list.append(est_dt)
        rds_1_3_cpu_list.append(c)
        
    rds_1_3_d = {"Dates : " : rds_1_3_date_list, "FreeableMemory(in GiB) : " : rds_1_3_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    rds_1_3_df = pd.DataFrame(rds_1_3_d, columns=['Dates : ', 'FreeableMemory(in GiB) : '])
    rds_1_3_df.sort_values(by=['Dates : ','FreeableMemory(in GiB) : '], inplace = True)
    rds_1_3_Timestamp = rds_1_3_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    rds_1_3_Timestamp_list = []
    
    for j in rds_1_3_Timestamp:
        time = j.split()
        time = time[1]
        rds_1_3_Timestamp_list.append(time)
    rds_1_3_CPU_Utilization = rds_1_3_df['FreeableMemory(in GiB) : '].tolist()
    
    fig3 = plt.rcParams['font.size'] = '7'
    
    plt.plot(rds_1_3_Timestamp_list, rds_1_3_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Freeable Memory for RDS Aurora PostgresSQL(Writer)', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Freeable Memory', fontsize = 10)
    plt.grid()
    plt.savefig('rds_1_3_Freeable_Memory.png', dpi=100)
    plt.close()
    
    #-----------------RDS Writer Instance Write IOPS ---------------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    rds_1_4_response = client.get_metric_statistics(
        Namespace='AWS/RDS',
        MetricName='WriteIOPS',
        Dimensions=[
            {
                'Name': 'DBClusterIdentifier',
                'Value': 'prod-amc-db-cluster'
            },
            
            {
                'Name':'Role',
                'Value':'WRITER'
            },
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    
    rds_1_4_a = (rds_1_4_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    rds_1_4_date_list = []
    rds_1_4_cpu_list = []
    for i in rds_1_4_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        rds_1_4_date_list.append(est_dt)
        rds_1_4_cpu_list.append(c)
        
    rds_1_4_d = {"Dates : " : rds_1_4_date_list, "Write IOPS : " : rds_1_4_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    rds_1_4_df = pd.DataFrame(rds_1_4_d, columns=['Dates : ', 'Write IOPS : '])
    rds_1_4_df.sort_values(by=['Dates : ','Write IOPS : '], inplace = True)
    rds_1_4_Timestamp = rds_1_4_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    rds_1_4_Timestamp_list = []
    
    for j in rds_1_4_Timestamp:
        time = j.split()
        time = time[1]
        rds_1_4_Timestamp_list.append(time)
    rds_1_4_CPU_Utilization = rds_1_4_df['Write IOPS : '].tolist()
    
    fig4 = plt.rcParams['font.size'] = '7'
    
    plt.plot(rds_1_4_Timestamp_list, rds_1_4_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Write IOPS for RDS Aurora PostgresSQL(Writer)', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Write IOPS', fontsize = 10)
    plt.grid()
    plt.savefig('rds_1_4_Write_IOPS.png', dpi=100)
    plt.close()
    
    
    #-----------------RDS Writer Instance Reader IOPS---------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    rds_1_5_response = client.get_metric_statistics(
        Namespace='AWS/RDS',
        MetricName='ReadIOPS',
        Dimensions=[
            {
                'Name': 'DBClusterIdentifier',
                'Value': 'prod-amc-db-cluster'
            },
            {
                'Name':'Role',
                'Value':'WRITER'
            },
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    
    rds_1_5_a = (rds_1_5_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    rds_1_5_date_list = []
    rds_1_5_cpu_list = []
    for i in rds_1_5_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        rds_1_5_date_list.append(est_dt)
        rds_1_5_cpu_list.append(c)
        
    rds_1_5_d = {"Dates : " : rds_1_5_date_list, "Read IOPS : " : rds_1_5_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    rds_1_5_df = pd.DataFrame(rds_1_5_d, columns=['Dates : ', 'Read IOPS : '])
    rds_1_5_df.sort_values(by=['Dates : ','Read IOPS : '], inplace = True)
    rds_1_5_Timestamp = rds_1_5_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    rds_1_5_Timestamp_list = []
    
    for j in rds_1_5_Timestamp:
        time = j.split()
        time = time[1]
        rds_1_5_Timestamp_list.append(time)
    rds_1_5_CPU_Utilization = rds_1_5_df['Read IOPS : '].tolist()
    
    fig5 = plt.rcParams['font.size'] = '7'
    
    plt.plot(rds_1_5_Timestamp_list, rds_1_5_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Reader IOPS for RDS Aurora PostgresSQL(Writer)', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Writer IOPS', fontsize = 10)
    plt.grid()
    plt.savefig('rds_1_5_Reader_IOPS.png', dpi=100)
    plt.close()
    
    #------------------ RDS Aurora Writer Instance Queue Depth ---------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    rds_1_6_response = client.get_metric_statistics(
        Namespace='AWS/RDS',
        MetricName='DiskQueueDepth',
        Dimensions=[
            {
                'Name': 'DBClusterIdentifier',
                'Value': 'prod-amc-db-cluster'
            },
            {
                'Name':'Role',
                'Value':'WRITER'
            },
        ],
        StartTime=datetime(2022,6,15,22,0),
        EndTime=datetime(2022,6,16,10,30),
        #StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        #EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    
    rds_1_6_a = (rds_1_6_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    rds_1_6_date_list = []
    rds_1_6_cpu_list = []
    for i in rds_1_6_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        rds_1_6_date_list.append(est_dt)
        rds_1_6_cpu_list.append(c)
        
    rds_1_6_d = {"Dates : " : rds_1_6_date_list, "Queue Depth : " : rds_1_6_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    rds_1_6_df = pd.DataFrame(rds_1_6_d, columns=['Dates : ', 'Queue Depth : '])
    rds_1_6_df.sort_values(by=['Dates : ','Queue Depth : '], inplace = True)
    rds_1_6_Timestamp = rds_1_6_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    rds_1_6_Timestamp_list = []
    
    for j in rds_1_6_Timestamp:
        time = j.split()
        time = time[1]
        rds_1_6_Timestamp_list.append(time)
    rds_1_6_CPU_Utilization = rds_1_6_df['Queue Depth : '].tolist()
    
    fig6 = plt.rcParams['font.size'] = '7'
    
    plt.plot(rds_1_6_Timestamp_list, rds_1_6_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Queue Depth for RDS Aurora PostgresSQL(Writer)', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Queue Depth', fontsize = 10)
    plt.grid()
    plt.savefig('rds_1_6_Queue_Depth.png', dpi=100)
    plt.close()
    
    #---------------RDS Aurora Reader Instance CPU Utilization-----------------------
    
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    rds_2_1_response = client.get_metric_statistics(
        Namespace='AWS/RDS',
        MetricName='CPUUtilization',
        Dimensions=[
            {
                'Name': 'DBClusterIdentifier',
                'Value': 'prod-amc-db-cluster',
            },
            {
                'Name':'Role',
                'Value':'READER'
            },
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    
    rds_2_1_a = (rds_2_1_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    rds_2_1_date_list = []
    rds_2_1_cpu_list = []
    for i in rds_2_1_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        rds_2_1_date_list.append(est_dt)
        rds_2_1_cpu_list.append(c)
        
    rds_2_1_d = {"Dates : " : rds_2_1_date_list, "CPU : " : rds_2_1_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    rds_2_1_df = pd.DataFrame(rds_2_1_d, columns=['Dates : ', 'CPU : '])
    rds_2_1_df.sort_values(by=['Dates : ','CPU : '], inplace = True)
    rds_2_1_Timestamp = rds_2_1_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    rds_2_1_Timestamp_list = []
    
    for j in rds_2_1_Timestamp:
        time = j.split()
        time = time[1]
        rds_2_1_Timestamp_list.append(time)
    rds_2_1_CPU_Utilization = rds_2_1_df['CPU : '].tolist()
    
    fig7 = plt.rcParams['font.size'] = '7'
    
    plt.plot(rds_2_1_Timestamp_list, rds_2_1_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('CPU Utilization of RDS Aurora PostgresSQL(Reader)', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('CPU Utilization', fontsize = 10)
    plt.grid()
    plt.savefig('rds_2_1_CPU_Utilization.png', dpi=100)
    plt.close()
    
    #------------------RDS Aurora Reader Instance DB Connections ----------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    rds_2_2_response = client.get_metric_statistics(
        Namespace='AWS/RDS',
        MetricName='DatabaseConnections',
        Dimensions=[
            {
                'Name': 'DBClusterIdentifier',
                'Value': 'prod-amc-db-cluster'
            },
            {
                'Name':'Role',
                'Value':'READER'
            },
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    
    rds_2_2_a = (rds_2_2_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    rds_2_2_date_list = []
    rds_2_2_cpu_list = []
    for i in rds_2_2_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        rds_2_2_date_list.append(est_dt)
        rds_2_2_cpu_list.append(c)
        
    rds_2_2_d = {"Dates : " : rds_2_2_date_list, "DB Connections : " : rds_2_2_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    rds_2_2_df = pd.DataFrame(rds_2_2_d, columns=['Dates : ', 'DB Connections : '])
    rds_2_2_df.sort_values(by=['Dates : ','DB Connections : '], inplace = True)
    rds_2_2_Timestamp = rds_2_2_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    rds_2_2_Timestamp_list = []
    
    for j in rds_2_2_Timestamp:
        time = j.split()
        time = time[1]
        rds_2_2_Timestamp_list.append(time)
    rds_2_2_CPU_Utilization = rds_2_2_df['DB Connections : '].tolist()
    
    fig8 = plt.rcParams['font.size'] = '7'
    
    plt.plot(rds_2_2_Timestamp_list, rds_2_2_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('DB Connections for RDS Aurora PostgresSQL(Reader)', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('DB Connections', fontsize = 10)
    plt.grid()
    plt.savefig('rds_2_2_DB_Connections.png', dpi=100)
    plt.close()
    
    #---------------RDS Aurora Reader Instance Freeable Memory ----------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    rds_2_3_response = client.get_metric_statistics(
        Namespace='AWS/RDS',
        MetricName='FreeableMemory',
        Dimensions=[
            {
                'Name': 'DBClusterIdentifier',
                'Value': 'prod-amc-db-cluster'
            },
            {
                'Name':'Role',
                'Value':'READER'
            },
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    
    rds_2_3_a = (rds_2_3_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    rds_2_3_date_list = []
    rds_2_3_cpu_list = []
    for i in rds_2_3_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        rds_2_3_date_list.append(est_dt)
        rds_2_3_cpu_list.append(c)
        
    rds_2_3_d = {"Dates : " : rds_2_3_date_list, "FreeableMemory(in GiB) : " : rds_2_3_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    rds_2_3_df = pd.DataFrame(rds_2_3_d, columns=['Dates : ', 'FreeableMemory(in GiB) : '])
    rds_2_3_df.sort_values(by=['Dates : ','FreeableMemory(in GiB) : '], inplace = True)
    rds_2_3_Timestamp = rds_2_3_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    rds_2_3_Timestamp_list = []
    
    for j in rds_2_3_Timestamp:
        time = j.split()
        time = time[1]
        rds_2_3_Timestamp_list.append(time)
    rds_2_3_CPU_Utilization = rds_2_3_df['FreeableMemory(in GiB) : '].tolist()
    
    fig9 = plt.rcParams['font.size'] = '7'
    
    plt.plot(rds_2_3_Timestamp_list, rds_2_3_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Freeable Memory for RDS Aurora PostgresSQL(Writer)', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Freeable Memory', fontsize = 10)
    plt.grid(True)
    plt.savefig('rds_2_3_Freeable_Memory.png', dpi=100)
    plt.close()
    
    #-------------RDS Aurora Reader Instance Writer IOPS ------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    rds_2_4_response = client.get_metric_statistics(
        Namespace='AWS/RDS',
        MetricName='WriteIOPS',
        Dimensions=[
            {
                'Name': 'DBClusterIdentifier',
                'Value': 'prod-amc-db-cluster'
            },
            {
                'Name':'Role',
                'Value':'READER'
            },
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    
    rds_2_4_a = (rds_2_4_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    rds_2_4_date_list = []
    rds_2_4_cpu_list = []
    for i in rds_2_4_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        rds_2_4_date_list.append(est_dt)
        rds_2_4_cpu_list.append(c)
        
    rds_2_4_d = {"Dates : " : rds_2_4_date_list, "Write IOPS : " : rds_2_4_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    rds_2_4_df = pd.DataFrame(rds_2_4_d, columns=['Dates : ', 'Write IOPS : '])
    rds_2_4_df.sort_values(by=['Dates : ','Write IOPS : '], inplace = True)
    rds_2_4_Timestamp = rds_2_4_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    rds_2_4_Timestamp_list = []
    
    for j in rds_2_4_Timestamp:
        time = j.split()
        time = time[1]
        rds_2_4_Timestamp_list.append(time)
    rds_2_4_CPU_Utilization = rds_2_4_df['Write IOPS : '].tolist()
    
    fig10 = plt.rcParams['font.size'] = '7'
    
    plt.plot(rds_2_4_Timestamp_list, rds_2_4_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Write IOPS for RDS Aurora PostgresSQL(Reader)', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Write IOPS', fontsize = 10)
    plt.grid(True)
    plt.savefig('rds_2_4_Write_IOPS.png', dpi=100)
    plt.close()
    
    #----------------RDS Aurora Reader Instance Reader IOPS ---------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    rds_2_5_response = client.get_metric_statistics(
        Namespace='AWS/RDS',
        MetricName='ReadIOPS',
        Dimensions=[
            {
                'Name': 'DBClusterIdentifier',
                'Value': 'prod-amc-db-cluster'
            },
            {
                'Name':'Role',
                'Value':'READER'
            },
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    
    rds_2_5_a = (rds_2_5_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    rds_2_5_date_list = []
    rds_2_5_cpu_list = []
    for i in rds_2_5_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        rds_2_5_date_list.append(est_dt)
        rds_2_5_cpu_list.append(c)
        
    rds_2_5_d = {"Dates : " : rds_2_5_date_list, "Read IOPS : " : rds_2_5_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    rds_2_5_df = pd.DataFrame(rds_2_5_d, columns=['Dates : ', 'Read IOPS : '])
    rds_2_5_df.sort_values(by=['Dates : ','Read IOPS : '], inplace = True)
    rds_2_5_Timestamp = rds_2_5_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    rds_2_5_Timestamp_list = []
    
    for j in rds_2_5_Timestamp:
        time = j.split()
        time = time[1]
        rds_2_5_Timestamp_list.append(time)
    rds_2_5_CPU_Utilization = rds_2_5_df['Read IOPS : '].tolist()
    
    fig11 = plt.rcParams['font.size'] = '7'
    
    plt.plot(rds_2_5_Timestamp_list, rds_2_5_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Reader IOPS for RDS Aurora PostgresSQL(Reader)', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Reader IOPS', fontsize = 10)
    plt.grid(True)
    plt.savefig('rds_2_5_Reader_IOPS.png', dpi=100)
    plt.close()
    
    #---------------------- RDS Aurora Reader Instance Queue Depth --------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    rds_2_6_response = client.get_metric_statistics(
        Namespace='AWS/RDS',
        MetricName='DiskQueueDepth',
        Dimensions=[
            {
                'Name': 'DBClusterIdentifier',
                'Value': 'prod-amc-db-cluster'
            },
            {
                'Name':'Role',
                'Value':'READER'
            },
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    
    rds_2_6_a = (rds_2_6_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    rds_2_6_date_list = []
    rds_2_6_cpu_list = []
    for i in rds_2_6_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        rds_2_6_date_list.append(est_dt)
        rds_2_6_cpu_list.append(c)
        
    rds_2_6_d = {"Dates : " : rds_2_6_date_list, "Queue Depth : " : rds_2_6_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    rds_2_6_df = pd.DataFrame(rds_2_6_d, columns=['Dates : ', 'Queue Depth : '])
    rds_2_6_df.sort_values(by=['Dates : ','Queue Depth : '], inplace = True)
    rds_2_6_Timestamp = rds_2_6_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    rds_2_6_Timestamp_list = []
    
    for j in rds_2_6_Timestamp:
        time = j.split()
        time = time[1]
        rds_2_6_Timestamp_list.append(time)
    rds_2_6_CPU_Utilization = rds_2_6_df['Queue Depth : '].tolist()
    
    fig12 = plt.rcParams['font.size'] = '7'
    
    plt.plot(rds_2_6_Timestamp_list, rds_2_6_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Queue Depth for RDS Aurora PostgresSQL(Reader)', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Queue Depth', fontsize = 10)
    plt.grid(True)
    plt.savefig('rds_2_6_Queue_Depth.png', dpi=100)
    plt.close()
    
    #-------------------Redshift Writer Instance CPU Utilization --------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    rs_1_1_response = client.get_metric_statistics(
        Namespace='AWS/Redshift',
        MetricName='CPUUtilization',
        Dimensions=[
            {
                'Name': 'ClusterIdentifier',
                'Value': 'americold-redshift-prod'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(rs_1_1_response)
    rs_1_1_a = (rs_1_1_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    rs_1_1_date_list = []
    rs_1_1_cpu_list = []
    for i in rs_1_1_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        rs_1_1_date_list.append(est_dt)
        rs_1_1_cpu_list.append(c)
        
    rs_1_1_d = {"Dates : " : rs_1_1_date_list, "CPU : " : rs_1_1_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    rs_1_1_df = pd.DataFrame(rs_1_1_d, columns=['Dates : ', 'CPU : '])
    rs_1_1_df.sort_values(by=['Dates : ','CPU : '], inplace = True)
    rs_1_1_Timestamp = rs_1_1_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    rs_1_1_Timestamp_list = []
    
    for j in rs_1_1_Timestamp:
        time = j.split()
        time = time[1]
        rs_1_1_Timestamp_list.append(time)
    rs_1_1_CPU_Utilization = rs_1_1_df['CPU : '].tolist()
    
    fig13 = plt.rcParams['font.size'] = '7'
    
    plt.plot(rs_1_1_Timestamp_list, rs_1_1_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('CPU Utilization for Reshift Writer Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('CPU Utilization', fontsize = 10)
    plt.grid(True)
    plt.savefig('rs_1_1_CPU_Utilization.png', dpi=100)
    plt.close()
    
    #-------------------Redshift Writer Insatance % Disk Space Used -------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    rs_1_2_response = client.get_metric_statistics(
        Namespace='AWS/Redshift',
        MetricName='PercentageDiskSpaceUsed',
        Dimensions=[
            {
                'Name': 'ClusterIdentifier',
                'Value': 'americold-redshift-prod'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(rs_1_2_response)
    rs_1_2_a = (rs_1_2_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    rs_1_2_date_list = []
    rs_1_2_cpu_list = []
    for i in rs_1_2_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        rs_1_2_date_list.append(est_dt)
        rs_1_2_cpu_list.append(c)
        
    rs_1_2_d = {"Dates : " : rs_1_2_date_list, "% Disk Space Used : " : rs_1_2_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    rs_1_2_df = pd.DataFrame(rs_1_2_d, columns=['Dates : ', '% Disk Space Used : '])
    rs_1_2_df.sort_values(by=['Dates : ','% Disk Space Used : '], inplace = True)
    rs_1_2_Timestamp = rs_1_2_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    rs_1_2_Timestamp_list = []
    
    for j in rs_1_2_Timestamp:
        time = j.split()
        time = time[1]
        rs_1_2_Timestamp_list.append(time)
    rs_1_2_CPU_Utilization = rs_1_2_df['% Disk Space Used : '].tolist()
    
    fig14 = plt.rcParams['font.size'] = '7'
    
    plt.plot(rs_1_2_Timestamp_list, rs_1_2_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('% of Diskspace used for Reshift Writer Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('% Disk Space Used', fontsize = 10)
    plt.grid(True)
    plt.savefig('rs_1_2_%diskspace_used.png', dpi=100)
    plt.close()
    
    #----------------Redshift Writer Instance Health Status --------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    rs_1_3_response = client.get_metric_statistics(
        Namespace='AWS/Redshift',
        MetricName='HealthStatus',
        Dimensions=[
            {
                'Name': 'ClusterIdentifier',
                'Value': 'americold-redshift-prod'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(rs_1_3_response)
    rs_1_3_a = (rs_1_3_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    rs_1_3_date_list = []
    rs_1_3_cpu_list = []
    for i in rs_1_3_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        rs_1_3_date_list.append(est_dt)
        rs_1_3_cpu_list.append(c)
        
    rs_1_3_d = {"Dates : " : rs_1_3_date_list, "Health Status : " : rs_1_3_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    rs_1_3_df = pd.DataFrame(rs_1_3_d, columns=['Dates : ', 'Health Status : '])
    rs_1_3_df.sort_values(by=['Dates : ','Health Status : '], inplace = True)
    rs_1_3_Timestamp = rs_1_3_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    rs_1_3_Timestamp_list = []
    
    for j in rs_1_3_Timestamp:
        time = j.split()
        time = time[1]
        rs_1_3_Timestamp_list.append(time)
    rs_1_3_CPU_Utilization = rs_1_3_df['Health Status : '].tolist()
    
    fig15 = plt.rcParams['font.size'] = '7'
    
    plt.plot(rs_1_3_Timestamp_list, rs_1_3_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Health Status for Reshift Writer Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Health Status', fontsize = 10)
    plt.grid(True)
    plt.savefig('rs_1_3_Health_Status.png', dpi=100)
    plt.close()
    
    #-----------------Redshift Reader Instance CPU Utilization -----------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    rs_2_1_response = client.get_metric_statistics(
        Namespace='AWS/Redshift',
        MetricName='CPUUtilization',
        Dimensions=[
            {
                'Name': 'ClusterIdentifier',
                'Value': 'americold-redshift-prod-readonly'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    rs_2_1_a = (rs_2_1_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    rs_2_1_date_list = []
    rs_2_1_cpu_list = []
    for i in rs_2_1_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        rs_2_1_date_list.append(est_dt)
        rs_2_1_cpu_list.append(c)
        
    rs_2_1_d = {"Dates : " : rs_2_1_date_list, "CPU : " : rs_2_1_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    rs_2_1_df = pd.DataFrame(rs_2_1_d, columns=['Dates : ', 'CPU : '])
    rs_2_1_df.sort_values(by=['Dates : ','CPU : '], inplace = True)
    rs_2_1_Timestamp = rs_2_1_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    rs_2_1_Timestamp_list = []
    
    for j in rs_2_1_Timestamp:
        time = j.split()
        time = time[1]
        rs_2_1_Timestamp_list.append(time)
    rs_2_1_CPU_Utilization = rs_2_1_df['CPU : '].tolist()
    
    
    fig16 = plt.rcParams['font.size'] = '7'
    
    plt.plot(rs_2_1_Timestamp_list, rs_2_1_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('CPU Utilization for Reshift Reader Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('CPU Utilization', fontsize = 10)
    plt.grid(True)
    plt.savefig('rs_2_1_CPU_Utilization.png', dpi=100)
    plt.close()
    
    #-----------------Redshift Reader Instance % Diskspace used ---------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    rs_2_2_response = client.get_metric_statistics(
        Namespace='AWS/Redshift',
        MetricName='PercentageDiskSpaceUsed',
        Dimensions=[
            {
                'Name': 'ClusterIdentifier',
                'Value': 'americold-redshift-prod-readonly'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(rs_2_2_response)
    rs_2_2_a = (rs_2_2_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    rs_2_2_date_list = []
    rs_2_2_cpu_list = []
    for i in rs_2_2_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        rs_2_2_date_list.append(est_dt)
        rs_2_2_cpu_list.append(c)
        
    rs_2_2_d = {"Dates : " : rs_2_2_date_list, "% Disk Space Used : " : rs_2_2_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    rs_2_2_df = pd.DataFrame(rs_2_2_d, columns=['Dates : ', '% Disk Space Used : '])
    rs_2_2_df.sort_values(by=['Dates : ','% Disk Space Used : '], inplace = True)
    rs_2_2_Timestamp = rs_2_2_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    rs_2_2_Timestamp_list = []
    
    for j in rs_2_2_Timestamp:
        time = j.split()
        time = time[1]
        rs_2_2_Timestamp_list.append(time)
    rs_2_2_CPU_Utilization = rs_2_2_df['% Disk Space Used : '].tolist()
    
    fig17 = plt.rcParams['font.size'] = '7'
    
    plt.plot(rs_2_2_Timestamp_list, rs_2_2_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('% of Diskspace used for Reshift Reader Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('% Disk Space Used', fontsize = 10)
    plt.grid(True)
    plt.savefig('rs_2_2_%diskspace_used.png', dpi=100)
    plt.close()
    
    #---------------Redshift Reader Instance Health Status -------------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    rs_2_3_response = client.get_metric_statistics(
        Namespace='AWS/Redshift',
        MetricName='HealthStatus',
        Dimensions=[
            {
                'Name': 'ClusterIdentifier',
                'Value': 'americold-redshift-prod-readonly'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(rs_2_3_response)
    rs_2_3_a = (rs_2_3_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    rs_2_3_date_list = []
    rs_2_3_cpu_list = []
    for i in rs_2_3_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        rs_2_3_date_list.append(est_dt)
        rs_2_3_cpu_list.append(c)
        
    rs_2_3_d = {"Dates : " : rs_2_3_date_list, "Health Status : " : rs_2_3_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    rs_2_3_df = pd.DataFrame(rs_2_3_d, columns=['Dates : ', 'Health Status : '])
    rs_2_3_df.sort_values(by=['Dates : ','Health Status : '], inplace = True)
    rs_2_3_Timestamp = rs_2_3_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    rs_2_3_Timestamp_list = []
    
    for j in rs_2_3_Timestamp:
        time = j.split()
        time = time[1]
        rs_2_3_Timestamp_list.append(time)
    rs_2_3_CPU_Utilization = rs_2_3_df['Health Status : '].tolist()
    
    fig18 = plt.rcParams['font.size'] = '7'
    
    plt.plot(rs_2_3_Timestamp_list, rs_2_3_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Health Status for Reshift Reader Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Health Status', fontsize = 10)
    plt.grid(True)
    plt.savefig('rs_2_3_Health_Status.png', dpi=100)
    plt.close()
    
    #---------------Storage Gateway Cache Percent Dirty ---------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    sg_3_response = client.get_metric_statistics(
        Namespace='AWS/StorageGateway',
        MetricName='CachePercentDirty',
        Dimensions=[
            {
                'Name': 'GatewayName',
                'Value': 'Americold_Storage_Gateway'
            },
            {
                'Name': 'GatewayId',
                'Value': 'sgw-E739D78E'
            },
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    
    sg_3_a = (sg_3_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    sg_3_date_list = []
    sg_3_cpu_list = []
    for i in sg_3_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        sg_3_date_list.append(est_dt)
        sg_3_cpu_list.append(c)
        
    sg_3_d = {"Dates : " : sg_3_date_list, "Cache Percent Dirty : " : sg_3_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    sg_3_df = pd.DataFrame(sg_3_d, columns=['Dates : ', 'Cache Percent Dirty : '])
    sg_3_df.sort_values(by=['Dates : ','Cache Percent Dirty : '], inplace = True)
    sg_3_Timestamp = sg_3_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    sg_3_Timestamp_list = []
    
    for j in sg_3_Timestamp:
        time = j.split()
        time = time[1]
        sg_3_Timestamp_list.append(time)
    sg_3_CPU_Utilization = sg_3_df['Cache Percent Dirty : '].tolist()
    
    fig19 = plt.rcParams['font.size'] = '7'
    
    plt.plot(sg_3_Timestamp_list, sg_3_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Cache Percent Dirty for Storage Gateway Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Cache Percent Dirty', fontsize = 10)
    plt.grid(True)
    plt.savefig('sg_3_Cache_Percent_Dirty.png', dpi=100)
    plt.close()
    
    #------------Storage Gateway Cache Hit Percent --------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    sg_4_response = client.get_metric_statistics(
        Namespace='AWS/StorageGateway',
        MetricName='CacheHitPercent',
        Dimensions=[
            {
                'Name': 'GatewayName',
                'Value': 'Americold_Storage_Gateway'
            },
            {
                'Name': 'GatewayId',
                'Value': 'sgw-E739D78E'
            },
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    
    sg_4_a = (sg_4_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    sg_4_date_list = []
    sg_4_cpu_list = []
    for i in sg_4_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        sg_4_date_list.append(est_dt)
        sg_4_cpu_list.append(c)
        
    sg_4_d = {"Dates : " : sg_4_date_list, "Cache Hit Percent : " : sg_4_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    sg_4_df = pd.DataFrame(sg_4_d, columns=['Dates : ', 'Cache Hit Percent : '])
    sg_4_df.sort_values(by=['Dates : ','Cache Hit Percent : '], inplace = True)
    sg_4_Timestamp = sg_4_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    sg_4_Timestamp_list = []
    
    for j in sg_4_Timestamp:
        time = j.split()
        time = time[1]
        sg_4_Timestamp_list.append(time)
    sg_4_CPU_Utilization = sg_4_df['Cache Hit Percent : '].tolist()
    
    fig20 = plt.rcParams['font.size'] = '7'
    
    fig = plt.plot(sg_4_Timestamp_list, sg_4_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Cache Hit Dirty for Storage Gateway Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Cache Hit Dirty', fontsize = 10)
    plt.grid(True)
    plt.savefig('sg_4_Cache_Hit_Dirty.png', dpi=100)
    plt.close()
    
    #-------------------Document DB Writer Instance CPU Utilization -----------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    db_1_1_response = client.get_metric_statistics(
        Namespace='AWS/DocDB',
        MetricName='CPUUtilization',
        Dimensions=[
            {
                'Name': 'DBInstanceIdentifier',
                'Value': 'prod-i3pl-docdb3'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(db_1_1_response)
    db_1_1_a = (db_1_1_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    db_1_1_date_list = []
    db_1_1_cpu_list = []
    for i in db_1_1_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        db_1_1_date_list.append(est_dt)
        db_1_1_cpu_list.append(c)
        
    db_1_1_d = {"Dates : " : db_1_1_date_list, "CPU Utilization : " : db_1_1_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    db_1_1_df = pd.DataFrame(db_1_1_d, columns=['Dates : ', 'CPU Utilization : '])
    db_1_1_df.sort_values(by=['Dates : ','CPU Utilization : '], inplace = True)
    db_1_1_Timestamp = db_1_1_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    db_1_1_Timestamp_list = []
    
    for j in db_1_1_Timestamp:
        time = j.split()
        time = time[1]
        db_1_1_Timestamp_list.append(time)
    db_1_1_CPU_Utilization = db_1_1_df['CPU Utilization : '].tolist()
    
    fig21 = plt.rcParams['font.size'] = '7'
    
    plt.plot(db_1_1_Timestamp_list, db_1_1_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('CPU Utilization for Document DB Writer Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('CPU Utilization', fontsize = 10)
    plt.grid(True)
    plt.savefig('db_1_1_CPU_Utilization.png', dpi=100)
    plt.close()
    
    #------------Document DB Writer Instance Freeable Memory --------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    db_1_2_response = client.get_metric_statistics(
        Namespace='AWS/DocDB',
        MetricName='FreeableMemory',
        Dimensions=[
            {
                'Name': 'DBInstanceIdentifier',
                'Value': 'prod-i3pl-docdb3'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(db_1_2_response)
    db_1_2_a = (db_1_2_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    db_1_2_date_list = []
    db_1_2_cpu_list = []
    for i in db_1_2_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        db_1_2_date_list.append(est_dt)
        db_1_2_cpu_list.append(c)
        
    db_1_2_d = {"Dates : " : db_1_2_date_list, "Freeable Memory : " : db_1_2_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    db_1_2_df = pd.DataFrame(db_1_2_d, columns=['Dates : ', 'Freeable Memory : '])
    db_1_2_df.sort_values(by=['Dates : ','Freeable Memory : '], inplace = True)
    db_1_2_Timestamp = db_1_2_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    db_1_2_Timestamp_list = []
    
    for j in db_1_2_Timestamp:
        time = j.split()
        time = time[1]
        db_1_2_Timestamp_list.append(time)
    db_1_2_CPU_Utilization = db_1_2_df['Freeable Memory : '].tolist()
    
    fig22 = plt.rcParams['font.size'] = '7'
    
    plt.plot(db_1_2_Timestamp_list, db_1_2_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Freeable Memory for Document DB Writer Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Freeable Memory', fontsize = 10)
    plt.grid(True)
    plt.savefig('db_1_2_Freeable_Memory.png', dpi=100)
    plt.close()
    
    #---------------Document DB Writer Insatance Free Local Storage --------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    db_1_3_response = client.get_metric_statistics(
        Namespace='AWS/DocDB',
        MetricName='FreeLocalStorage',
        Dimensions=[
            {
                'Name': 'DBInstanceIdentifier',
                'Value': 'prod-i3pl-docdb3'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(db_1_3_response)
    db_1_3_a = (db_1_3_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    db_1_3_date_list = []
    db_1_3_cpu_list = []
    for i in db_1_3_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        db_1_3_date_list.append(est_dt)
        db_1_3_cpu_list.append(c)
        
    db_1_3_d = {"Dates : " : db_1_3_date_list, "Free Local Storage : " : db_1_3_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    db_1_3_df = pd.DataFrame(db_1_3_d, columns=['Dates : ', 'Free Local Storage : '])
    db_1_3_df.sort_values(by=['Dates : ','Free Local Storage : '], inplace = True)
    db_1_3_Timestamp = db_1_3_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    db_1_3_Timestamp_list = []
    
    for j in db_1_3_Timestamp:
        time = j.split()
        time = time[1]
        db_1_3_Timestamp_list.append(time)
    db_1_3_CPU_Utilization = db_1_3_df['Free Local Storage : '].tolist()
    
    fig23 = plt.rcParams['font.size'] = '7'
    
    plt.plot(db_1_3_Timestamp_list, db_1_3_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Free Local Storage for Document DB Writer Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Free Local Storage', fontsize = 10)
    plt.grid(True)
    plt.savefig('db_1_3_Free_Local_Storage.png', dpi=100)
    plt.close()
    
    #-------------Document DB Writer Instance Swap Usage ---------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    db_1_4_response = client.get_metric_statistics(
        Namespace='AWS/DocDB',
        MetricName='SwapUsage',
        Dimensions=[
            {
                'Name': 'DBInstanceIdentifier',
                'Value': 'prod-i3pl-docdb3'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(db_1_4_response)
    db_1_4_a = (db_1_4_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    db_1_4_date_list = []
    db_1_4_cpu_list = []
    for i in db_1_4_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        db_1_4_date_list.append(est_dt)
        db_1_4_cpu_list.append(c)
        
    db_1_4_d = {"Dates : " : db_1_4_date_list, "Swap Usage : " : db_1_4_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    db_1_4_df = pd.DataFrame(db_1_4_d, columns=['Dates : ', 'Swap Usage : '])
    db_1_4_df.sort_values(by=['Dates : ','Swap Usage : '], inplace = True)
    db_1_4_Timestamp = db_1_4_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    db_1_4_Timestamp_list = []
    
    for j in db_1_4_Timestamp:
        time = j.split()
        time = time[1]
        db_1_4_Timestamp_list.append(time)
    db_1_4_CPU_Utilization = db_1_4_df['Swap Usage : '].tolist()
    
    fig24 = plt.rcParams['font.size'] = '7'
    
    plt.plot(db_1_4_Timestamp_list, db_1_4_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Swap Usage for Document DB Writer Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Swap Usage', fontsize = 10)
    plt.grid(True)
    plt.savefig('db_1_4_Swap_Usage.png', dpi=100)
    plt.close()
    
    #-----------------Document DB Database Connections ------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    db_1_5_response = client.get_metric_statistics(
        Namespace='AWS/DocDB',
        MetricName='DatabaseConnections',
        Dimensions=[
            {
                'Name': 'DBInstanceIdentifier',
                'Value': 'prod-i3pl-docdb3'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(db_1_5_response)
    db_1_5_a = (db_1_5_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    db_1_5_date_list = []
    db_1_5_cpu_list = []
    for i in db_1_5_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        db_1_5_date_list.append(est_dt)
        db_1_5_cpu_list.append(c)
        
    db_1_5_d = {"Dates : " : db_1_5_date_list, "Database Connections : " : db_1_5_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    db_1_5_df = pd.DataFrame(db_1_5_d, columns=['Dates : ', 'Database Connections : '])
    db_1_5_df.sort_values(by=['Dates : ','Database Connections : '], inplace = True)
    db_1_5_Timestamp = db_1_5_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    db_1_5_Timestamp_list = []
    
    for j in db_1_5_Timestamp:
        time = j.split()
        time = time[1]
        db_1_5_Timestamp_list.append(time)
    db_1_5_CPU_Utilization = db_1_5_df['Database Connections : '].tolist()
    
    fig26 = plt.rcParams['font.size'] = '7'
    
    plt.plot(db_1_5_Timestamp_list, db_1_5_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Database Connections for Document DB Writer Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Database Connections', fontsize = 10)
    plt.grid(True)
    plt.savefig('db_1_5_Database_Connections.png', dpi=100)
    plt.close()
    
    #---------------Document DB Database Connections Max ------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    db_1_6_response = client.get_metric_statistics(
        Namespace='AWS/DocDB',
        MetricName='DatabaseConnectionsMax',
        Dimensions=[
            {
                'Name': 'DBInstanceIdentifier',
                'Value': 'prod-i3pl-docdb3'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(db_1_6_response)
    db_1_6_a = (db_1_6_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    db_1_6_date_list = []
    db_1_6_cpu_list = []
    for i in db_1_6_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        db_1_6_date_list.append(est_dt)
        db_1_6_cpu_list.append(c)
        
    db_1_6_d = {"Dates : " : db_1_6_date_list, "Database Connections Max : " : db_1_6_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    db_1_6_df = pd.DataFrame(db_1_6_d, columns=['Dates : ', 'Database Connections Max : '])
    db_1_6_df.sort_values(by=['Dates : ','Database Connections Max : '], inplace = True)
    db_1_6_Timestamp = db_1_6_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    db_1_6_Timestamp_list = []
    
    for j in db_1_6_Timestamp:
        time = j.split()
        time = time[1]
        db_1_6_Timestamp_list.append(time)
    db_1_6_CPU_Utilization = db_1_6_df['Database Connections Max : '].tolist()
    
    fig26 = plt.rcParams['font.size'] = '7'
    
    plt.plot(db_1_6_Timestamp_list, db_1_6_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Maximum Database Connections for Document DB Writer Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Maximum Database Connections', fontsize = 10)
    plt.grid(True)
    plt.savefig('db_1_6_Maximum_Database_Connections.png', dpi=100)
    plt.close()
    
    #----------------Document DB Writer Instance Database Cursors --------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    db_1_7_response = client.get_metric_statistics(
        Namespace='AWS/DocDB',
        MetricName='DatabaseCursors',
        Dimensions=[
            {
                'Name': 'DBInstanceIdentifier',
                'Value': 'prod-i3pl-docdb3'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(db_1_7_response)
    db_1_7_a = (db_1_7_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    db_1_7_date_list = []
    db_1_7_cpu_list = []
    for i in db_1_7_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        db_1_7_date_list.append(est_dt)
        db_1_7_cpu_list.append(c)
        
    db_1_7_d = {"Dates : " : db_1_7_date_list, "Database Cursors : " : db_1_7_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    db_1_7_df = pd.DataFrame(db_1_7_d, columns=['Dates : ', 'Database Cursors : '])
    db_1_7_df.sort_values(by=['Dates : ','Database Cursors : '], inplace = True)
    db_1_7_Timestamp = db_1_7_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    db_1_7_Timestamp_list = []
    
    for j in db_1_7_Timestamp:
        time = j.split()
        time = time[1]
        db_1_7_Timestamp_list.append(time)
    db_1_7_CPU_Utilization = db_1_7_df['Database Cursors : '].tolist()
    
    fig27 = plt.rcParams['font.size'] = '7'
    
    plt.plot(db_1_7_Timestamp_list, db_1_7_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Database Cursors for Document DB Writer Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Database Cursors', fontsize = 10)
    plt.grid(True)
    plt.savefig('db_1_7_Database_Cursors.png', dpi=100)
    plt.close()
    
    #-------------Document DB Writer Instance Maximum Database Cursors --------------
    
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    db_1_8_response = client.get_metric_statistics(
        Namespace='AWS/DocDB',
        MetricName='DatabaseCursorsMax',
        Dimensions=[
            {
                'Name': 'DBInstanceIdentifier',
                'Value': 'prod-i3pl-docdb3'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(db_1_7_response)
    db_1_8_a = (db_1_7_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    db_1_8_date_list = []
    db_1_8_cpu_list = []
    for i in db_1_8_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        db_1_8_date_list.append(est_dt)
        db_1_8_cpu_list.append(c)
        
    db_1_8_d = {"Dates : " : db_1_8_date_list, "Database Cursors Max : " : db_1_8_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    db_1_8_df = pd.DataFrame(db_1_8_d, columns=['Dates : ', 'Database Cursors Max : '])
    db_1_8_df.sort_values(by=['Dates : ','Database Cursors Max : '], inplace = True)
    db_1_8_Timestamp = db_1_8_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    db_1_8_Timestamp_list = []
    
    for j in db_1_8_Timestamp:
        time = j.split()
        time = time[1]
        db_1_8_Timestamp_list.append(time)
    db_1_8_CPU_Utilization = db_1_8_df['Database Cursors Max : '].tolist()
    
    fig28 = plt.rcParams['font.size'] = '7'
    
    plt.plot(db_1_8_Timestamp_list, db_1_8_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Maximum Database Cursors for Document DB Writer Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Maximum Database Cursors', fontsize = 10)
    plt.grid(True)
    plt.savefig('db_1_8_Maximum_Database_Cursors.png', dpi=100)
    plt.close()
    
    #-------------Document DB Reader Insatances --------------------
    
    #-------------------Document DB Reader Instance CPU Utilization -----------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    db_2_1_response = client.get_metric_statistics(
        Namespace='AWS/DocDB',
        MetricName='CPUUtilization',
        Dimensions=[
            {
                'Name': 'DBInstanceIdentifier',
                'Value': 'prod-i3pl-docdb'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(db_1_1_response)
    db_2_1_a = (db_2_1_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    db_2_1_date_list = []
    db_2_1_cpu_list = []
    for i in db_2_1_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        db_2_1_date_list.append(est_dt)
        db_2_1_cpu_list.append(c)
        
    db_2_1_d = {"Dates : " : db_2_1_date_list, "CPU Utilization : " : db_2_1_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    db_2_1_df = pd.DataFrame(db_2_1_d, columns=['Dates : ', 'CPU Utilization : '])
    db_2_1_df.sort_values(by=['Dates : ','CPU Utilization : '], inplace = True)
    db_2_1_Timestamp = db_2_1_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    db_2_1_Timestamp_list = []
    
    for j in db_2_1_Timestamp:
        time = j.split()
        time = time[1]
        db_2_1_Timestamp_list.append(time)
    db_2_1_CPU_Utilization = db_2_1_df['CPU Utilization : '].tolist()
    
    fig29 = plt.rcParams['font.size'] = '7'
    
    plt.plot(db_2_1_Timestamp_list, db_2_1_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('CPU Utilization for Document DB Reader Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('CPU Utilization', fontsize = 10)
    plt.grid(True)
    plt.savefig('db_2_1_CPU_Utilization.png', dpi=100)
    plt.close()
    
    #------------Document DB Reader Instance Freeable Memory --------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    db_2_2_response = client.get_metric_statistics(
        Namespace='AWS/DocDB',
        MetricName='FreeableMemory',
        Dimensions=[
            {
                'Name': 'DBInstanceIdentifier',
                'Value': 'prod-i3pl-docdb'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(db_1_2_response)
    db_2_2_a = (db_2_2_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    db_2_2_date_list = []
    db_2_2_cpu_list = []
    for i in db_2_2_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        db_2_2_date_list.append(est_dt)
        db_2_2_cpu_list.append(c)
        
    db_2_2_d = {"Dates : " : db_2_2_date_list, "Freeable Memory : " : db_2_2_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    db_2_2_df = pd.DataFrame(db_2_2_d, columns=['Dates : ', 'Freeable Memory : '])
    db_2_2_df.sort_values(by=['Dates : ','Freeable Memory : '], inplace = True)
    db_2_2_Timestamp = db_2_2_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    db_2_2_Timestamp_list = []
    
    for j in db_2_2_Timestamp:
        time = j.split()
        time = time[1]
        db_2_2_Timestamp_list.append(time)
    db_2_2_CPU_Utilization = db_2_2_df['Freeable Memory : '].tolist()
    
    fig30 = plt.rcParams['font.size'] = '7'
    
    plt.plot(db_2_2_Timestamp_list, db_2_2_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Freeable Memory for Document DB Reader Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Freeable Memory', fontsize = 10)
    plt.grid(True)
    plt.savefig('db_2_2_Freeable_Memory.png', dpi=100)
    plt.close()
    
    #---------------Document DB Reader Insatance Free Local Storage --------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    db_2_3_response = client.get_metric_statistics(
        Namespace='AWS/DocDB',
        MetricName='FreeLocalStorage',
        Dimensions=[
            {
                'Name': 'DBInstanceIdentifier',
                'Value': 'prod-i3pl-docdb'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(db_1_3_response)
    db_2_3_a = (db_1_3_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    db_2_3_date_list = []
    db_2_3_cpu_list = []
    for i in db_2_3_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        db_2_3_date_list.append(est_dt)
        db_2_3_cpu_list.append(c)
        
    db_2_3_d = {"Dates : " : db_2_3_date_list, "Free Local Storage : " : db_1_3_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    db_2_3_df = pd.DataFrame(db_2_3_d, columns=['Dates : ', 'Free Local Storage : '])
    db_2_3_df.sort_values(by=['Dates : ','Free Local Storage : '], inplace = True)
    db_2_3_Timestamp = db_2_3_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    db_2_3_Timestamp_list = []
    
    for j in db_2_3_Timestamp:
        time = j.split()
        time = time[1]
        db_2_3_Timestamp_list.append(time)
    db_2_3_CPU_Utilization = db_2_3_df['Free Local Storage : '].tolist()
    
    fig31 = plt.rcParams['font.size'] = '7'
    
    plt.plot(db_2_3_Timestamp_list, db_2_3_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Free Local Storage for Document DB Reader Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Free Local Storage', fontsize = 10)
    plt.grid(True)
    plt.savefig('db_2_3_Free_Local_Storage.png', dpi=100)
    plt.close()
    
    #-------------Document DB Reader Instance Swap Usage ---------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    db_2_4_response = client.get_metric_statistics(
        Namespace='AWS/DocDB',
        MetricName='SwapUsage',
        Dimensions=[
            {
                'Name': 'DBInstanceIdentifier',
                'Value': 'prod-i3pl-docdb'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(db_1_4_response)
    db_2_4_a = (db_2_4_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    db_2_4_date_list = []
    db_2_4_cpu_list = []
    for i in db_2_4_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        db_2_4_date_list.append(est_dt)
        db_2_4_cpu_list.append(c)
        
    db_2_4_d = {"Dates : " : db_2_4_date_list, "Swap Usage : " : db_2_4_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    db_2_4_df = pd.DataFrame(db_2_4_d, columns=['Dates : ', 'Swap Usage : '])
    db_2_4_df.sort_values(by=['Dates : ','Swap Usage : '], inplace = True)
    db_2_4_Timestamp = db_2_4_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    db_2_4_Timestamp_list = []
    
    for j in db_2_4_Timestamp:
        time = j.split()
        time = time[1]
        db_2_4_Timestamp_list.append(time)
    db_2_4_CPU_Utilization = db_2_4_df['Swap Usage : '].tolist()
    
    fig32 = plt.rcParams['font.size'] = '7'
    
    plt.plot(db_2_4_Timestamp_list, db_2_4_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Swap Usage for Document DB Reader Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Swap Usage', fontsize = 10)
    plt.grid(True)
    plt.savefig('db_2_4_Swap_Usage.png', dpi=100)
    plt.close()
    
    #-----------------Document DB Reader Instance Database Connections ------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    db_2_5_response = client.get_metric_statistics(
        Namespace='AWS/DocDB',
        MetricName='DatabaseConnections',
        Dimensions=[
            {
                'Name': 'DBInstanceIdentifier',
                'Value': 'prod-i3pl-docdb'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(db_1_5_response)
    db_2_5_a = (db_2_5_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    db_2_5_date_list = []
    db_2_5_cpu_list = []
    for i in db_2_5_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        db_2_5_date_list.append(est_dt)
        db_2_5_cpu_list.append(c)
        
    db_2_5_d = {"Dates : " : db_2_5_date_list, "Database Connections : " : db_2_5_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    db_2_5_df = pd.DataFrame(db_2_5_d, columns=['Dates : ', 'Database Connections : '])
    db_2_5_df.sort_values(by=['Dates : ','Database Connections : '], inplace = True)
    db_2_5_Timestamp = db_2_5_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    db_2_5_Timestamp_list = []
    
    for j in db_2_5_Timestamp:
        time = j.split()
        time = time[1]
        db_2_5_Timestamp_list.append(time)
    db_2_5_CPU_Utilization = db_2_5_df['Database Connections : '].tolist()
    
    fig33 = plt.rcParams['font.size'] = '7'
    
    plt.plot(db_2_5_Timestamp_list, db_2_5_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Database Connections for Document DB Reader Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Database Connections', fontsize = 10)
    plt.grid(True)
    plt.savefig('db_2_5_Database_Connections.png', dpi=100)
    plt.close()
    
    #---------------Document DB Reader Instance Database Connections Max ------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    db_2_6_response = client.get_metric_statistics(
        Namespace='AWS/DocDB',
        MetricName='DatabaseConnectionsMax',
        Dimensions=[
            {
                'Name': 'DBInstanceIdentifier',
                'Value': 'prod-i3pl-docdb'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(db_1_6_response)
    db_2_6_a = (db_2_6_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    db_2_6_date_list = []
    db_2_6_cpu_list = []
    for i in db_2_6_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        db_2_6_date_list.append(est_dt)
        db_2_6_cpu_list.append(c)
        
    db_2_6_d = {"Dates : " : db_2_6_date_list, "Database Connections Max : " : db_2_6_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    db_2_6_df = pd.DataFrame(db_2_6_d, columns=['Dates : ', 'Database Connections Max : '])
    db_2_6_df.sort_values(by=['Dates : ','Database Connections Max : '], inplace = True)
    db_2_6_Timestamp = db_2_6_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    db_2_6_Timestamp_list = []
    
    for j in db_2_6_Timestamp:
        time = j.split()
        time = time[1]
        db_2_6_Timestamp_list.append(time)
    db_2_6_CPU_Utilization = db_2_6_df['Database Connections Max : '].tolist()
    
    fig34 = plt.rcParams['font.size'] = '7'
    
    plt.plot(db_2_6_Timestamp_list, db_2_6_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Maximum Database Connections for Document DB Reader Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Maximum Database Connections', fontsize = 10)
    plt.grid(True)
    plt.savefig('db_2_6_Maximum_Database_Connections.png', dpi=100)
    plt.close()
    
    #----------------Document DB Reader Instance Database Cursors --------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    db_2_7_response = client.get_metric_statistics(
        Namespace='AWS/DocDB',
        MetricName='DatabaseCursors',
        Dimensions=[
            {
                'Name': 'DBInstanceIdentifier',
                'Value': 'prod-i3pl-docdb'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(db_1_7_response)
    db_2_7_a = (db_2_7_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    db_2_7_date_list = []
    db_2_7_cpu_list = []
    for i in db_2_7_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        db_2_7_date_list.append(est_dt)
        db_2_7_cpu_list.append(c)
        
    db_2_7_d = {"Dates : " : db_2_7_date_list, "Database Cursors : " : db_2_7_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    db_2_7_df = pd.DataFrame(db_2_7_d, columns=['Dates : ', 'Database Cursors : '])
    db_2_7_df.sort_values(by=['Dates : ','Database Cursors : '], inplace = True)
    db_2_7_Timestamp = db_2_7_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    db_2_7_Timestamp_list = []
    
    for j in db_2_7_Timestamp:
        time = j.split()
        time = time[1]
        db_2_7_Timestamp_list.append(time)
    db_2_7_CPU_Utilization = db_2_7_df['Database Cursors : '].tolist()
    
    fig35 = plt.rcParams['font.size'] = '7'
    
    plt.plot(db_2_7_Timestamp_list, db_2_7_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Database Cursors for Document DB Reader Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Database Cursors', fontsize = 10)
    plt.grid(True)
    plt.savefig('db_2_7_Database_Cursors.png', dpi=100)
    plt.close()
    
    #-------------Document DB Reader Instance Maximum Database Cursors --------------
    
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    db_2_8_response = client.get_metric_statistics(
        Namespace='AWS/DocDB',
        MetricName='DatabaseCursorsMax',
        Dimensions=[
            {
                'Name': 'DBInstanceIdentifier',
                'Value': 'prod-i3pl-docdb'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(db_1_7_response)
    db_2_8_a = (db_2_8_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    db_2_8_date_list = []
    db_2_8_cpu_list = []
    for i in db_2_8_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        db_2_8_date_list.append(est_dt)
        db_2_8_cpu_list.append(c)
        
    db_2_8_d = {"Dates : " : db_2_8_date_list, "Database Cursors Max : " : db_2_8_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    db_2_8_df = pd.DataFrame(db_2_8_d, columns=['Dates : ', 'Database Cursors Max : '])
    db_2_8_df.sort_values(by=['Dates : ','Database Cursors Max : '], inplace = True)
    db_2_8_Timestamp = db_2_8_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    db_2_8_Timestamp_list = []
    
    for j in db_2_8_Timestamp:
        time = j.split()
        time = time[1]
        db_2_8_Timestamp_list.append(time)
    db_2_8_CPU_Utilization = db_2_8_df['Database Cursors Max : '].tolist()
    
    fig36 = plt.rcParams['font.size'] = '7'
    
    plt.plot(db_2_8_Timestamp_list, db_2_8_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Maximum Database Cursors for Document DB Reader Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Maximum Database Cursors', fontsize = 10)
    plt.grid(True)
    plt.savefig('db_2_8_Maximum_Database_Cursors.png', dpi=100)
    plt.close()
    
    #----------------- Storage Gateway EC2 Instance CPU Utilization-----------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    ec_1_1_response = client.get_metric_statistics(
        Namespace='AWS/EC2',
        MetricName='CPUUtilization',
        Dimensions=[
            {
                'Name': 'InstanceId',
                'Value': 'i-08db560943d4a48bd'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(ec_1_1_response)
    ec_1_1_a = (ec_1_1_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    ec_1_1_date_list = []
    ec_1_1_cpu_list = []
    for i in ec_1_1_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        ec_1_1_date_list.append(est_dt)
        ec_1_1_cpu_list.append(c)
        
    ec_1_1_d = {"Dates : " : ec_1_1_date_list, "CPU Utilization : " : ec_1_1_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    ec_1_1_df = pd.DataFrame(ec_1_1_d, columns=['Dates : ', 'CPU Utilization : '])
    ec_1_1_df.sort_values(by=['Dates : ','CPU Utilization : '], inplace = True)
    ec_1_1_Timestamp = ec_1_1_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    ec_1_1_Timestamp_list = []
    
    for j in ec_1_1_Timestamp:
        time = j.split()
        time = time[1]
        ec_1_1_Timestamp_list.append(time)
    ec_1_1_CPU_Utilization = ec_1_1_df['CPU Utilization : '].tolist()
    
    fig37 = plt.rcParams['font.size'] = '7'
    
    plt.plot(ec_1_1_Timestamp_list, ec_1_1_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('CPU Utilization for Storage Gateway EC2 Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('CPU Utilization', fontsize = 10)
    plt.grid(True)
    plt.savefig('ec_1_1_CPU_Utilization.png', dpi=100)
    plt.close()
    
    #------------------Storage Gateway EC2 Instance Status Check Failed(Any) ----------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    ec_1_2_response = client.get_metric_statistics(
        Namespace='AWS/EC2',
        MetricName='StatusCheckFailed',
        Dimensions=[
            {
                'Name': 'InstanceId',
                'Value': 'i-08db560943d4a48bd'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(ec_1_2_response)
    ec_1_2_a = (ec_1_2_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    ec_1_2_date_list = []
    ec_1_2_cpu_list = []
    for i in ec_1_2_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        ec_1_2_date_list.append(est_dt)
        ec_1_2_cpu_list.append(c)
        
    ec_1_2_d = {"Dates : " : ec_1_2_date_list, "Status Check Failed(Any) : " : ec_1_2_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    ec_1_2_df = pd.DataFrame(ec_1_2_d, columns=['Dates : ', 'Status Check Failed(Any) : '])
    ec_1_2_df.sort_values(by=['Dates : ','Status Check Failed(Any) : '], inplace = True)
    ec_1_2_Timestamp = ec_1_2_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    ec_1_2_Timestamp_list = []
    
    for j in ec_1_2_Timestamp:
        time = j.split()
        time = time[1]
        ec_1_2_Timestamp_list.append(time)
    ec_1_2_CPU_Utilization = ec_1_2_df['Status Check Failed(Any) : '].tolist()
    
    fig38 = plt.rcParams['font.size'] = '7'
    
    plt.plot(ec_1_2_Timestamp_list, ec_1_2_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Status Check Failed(Any) for Storage Gateway EC2 Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Status Check Failed(Any)', fontsize = 10)
    plt.grid(True)
    plt.savefig('ec_1_2_Status_Check_Failed(Any).png', dpi=100)
    plt.close()
    
    #-------------------Storage Gateway EC2 Insatance Status Check Failed(Instance) ------------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    ec_1_3_response = client.get_metric_statistics(
        Namespace='AWS/EC2',
        MetricName='StatusCheckFailed_Instance',
        Dimensions=[
            {
                'Name': 'InstanceId',
                'Value': 'i-08db560943d4a48bd'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(ec_1_3_response)
    ec_1_3_a = (ec_1_3_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    ec_1_3_date_list = []
    ec_1_3_cpu_list = []
    for i in ec_1_3_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        ec_1_3_date_list.append(est_dt)
        ec_1_3_cpu_list.append(c)
        
    ec_1_3_d = {"Dates : " : ec_1_3_date_list, "Status Check Failed(Instance) : " : ec_1_3_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    ec_1_3_df = pd.DataFrame(ec_1_3_d, columns=['Dates : ', 'Status Check Failed(Instance) : '])
    ec_1_3_df.sort_values(by=['Dates : ','Status Check Failed(Instance) : '], inplace = True)
    ec_1_3_Timestamp = ec_1_3_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    ec_1_3_Timestamp_list = []
    
    for j in ec_1_3_Timestamp:
        time = j.split()
        time = time[1]
        ec_1_3_Timestamp_list.append(time)
    ec_1_3_CPU_Utilization = ec_1_3_df['Status Check Failed(Instance) : '].tolist()
    
    fig39 = plt.rcParams['font.size'] = '7'
    
    plt.plot(ec_1_3_Timestamp_list, ec_1_3_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Status Check Failed(Instance) for Storage Gateway EC2 Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Status Check Failed(Instance)', fontsize = 10)
    plt.grid(True)
    plt.savefig('ec_1_3_Status_Check_Failed(Instance).png', dpi=100)
    plt.close()
    
    #---------------------Storage Gateway EC2 Instance Status_Check_Failed(System) -----------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    ec_1_4_response = client.get_metric_statistics(
        Namespace='AWS/EC2',
        MetricName='StatusCheckFailed_System',
        Dimensions=[
            {
                'Name': 'InstanceId',
                'Value': 'i-08db560943d4a48bd'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(ec_1_4_response)
    ec_1_4_a = (ec_1_4_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    ec_1_4_date_list = []
    ec_1_4_cpu_list = []
    for i in ec_1_4_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        ec_1_4_date_list.append(est_dt)
        ec_1_4_cpu_list.append(c)
        
    ec_1_4_d = {"Dates : " : ec_1_4_date_list, "Status Check Failed(System) : " : ec_1_4_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    ec_1_4_df = pd.DataFrame(ec_1_4_d, columns=['Dates : ', 'Status Check Failed(System) : '])
    ec_1_4_df.sort_values(by=['Dates : ','Status Check Failed(System) : '], inplace = True)
    ec_1_4_Timestamp = ec_1_4_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    ec_1_4_Timestamp_list = []
    
    for j in ec_1_4_Timestamp:
        time = j.split()
        time = time[1]
        ec_1_4_Timestamp_list.append(time)
    ec_1_4_CPU_Utilization = ec_1_4_df['Status Check Failed(System) : '].tolist()
    
    plt.rcParams['font.size'] = '7'
    
    fig40 = plt.plot(ec_1_4_Timestamp_list, ec_1_4_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Status Check Failed(System) for Storage Gateway EC2 Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Status Check Failed(System)', fontsize = 10)
    plt.grid(True)
    plt.savefig('ec_1_4_Status_Check_Failed(System).png', dpi=100)
    plt.close()
    
    #----------------PowerBI Gateway EC2 Instance CPU Utilization ------------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    ec_2_1_response = client.get_metric_statistics(
        Namespace='AWS/EC2',
        MetricName='CPUUtilization',
        Dimensions=[
            {
                'Name': 'InstanceId',
                'Value': 'i-0cc329b29763639a0'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(ec_2_1_response)
    ec_2_1_a = (ec_2_1_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    ec_2_1_date_list = []
    ec_2_1_cpu_list = []
    for i in ec_2_1_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        ec_2_1_date_list.append(est_dt)
        ec_2_1_cpu_list.append(c)
        
    ec_2_1_d = {"Dates : " : ec_2_1_date_list, "CPU Utilization : " : ec_2_1_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    ec_2_1_df = pd.DataFrame(ec_2_1_d, columns=['Dates : ', 'CPU Utilization : '])
    ec_2_1_df.sort_values(by=['Dates : ','CPU Utilization : '], inplace = True)
    ec_2_1_Timestamp = ec_2_1_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    ec_2_1_Timestamp_list = []
    
    for j in ec_2_1_Timestamp:
        time = j.split()
        time = time[1]
        ec_2_1_Timestamp_list.append(time)
    ec_2_1_CPU_Utilization = ec_2_1_df['CPU Utilization : '].tolist()
    
    plt.rcParams['font.size'] = '7'
    
    fig41 = plt.plot(ec_2_1_Timestamp_list, ec_2_1_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('CPU Utilization for PowerBI Gateway EC2 Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('CPU Utilization', fontsize = 10)
    plt.grid(True)
    plt.savefig('ec_2_1_CPU_Utilization.png', dpi=100)
    plt.close()
    
    #--------------------PowerBI Gateway EC2 Instance Status Check Failed(Any)----------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    ec_2_2_response = client.get_metric_statistics(
        Namespace='AWS/EC2',
        MetricName='StatusCheckFailed',
        Dimensions=[
            {
                'Name': 'InstanceId',
                'Value': 'i-0cc329b29763639a0'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(ec_1_2_response)
    ec_2_2_a = (ec_2_2_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    ec_2_2_date_list = []
    ec_2_2_cpu_list = []
    for i in ec_2_2_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        ec_2_2_date_list.append(est_dt)
        ec_2_2_cpu_list.append(c)
        
    ec_2_2_d = {"Dates : " : ec_2_2_date_list, "Status Check Failed(Any) : " : ec_2_2_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    ec_2_2_df = pd.DataFrame(ec_2_2_d, columns=['Dates : ', 'Status Check Failed(Any) : '])
    ec_2_2_df.sort_values(by=['Dates : ','Status Check Failed(Any) : '], inplace = True)
    ec_2_2_Timestamp = ec_2_2_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    ec_2_2_Timestamp_list = []
    
    for j in ec_2_2_Timestamp:
        time = j.split()
        time = time[1]
        ec_2_2_Timestamp_list.append(time)
    ec_2_2_CPU_Utilization = ec_2_2_df['Status Check Failed(Any) : '].tolist()
    
    plt.rcParams['font.size'] = '7'
    
    fig42 = plt.plot(ec_2_2_Timestamp_list, ec_2_2_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Status Check Failed(Any) for PowerBI Gateway EC2 Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Status Check Failed(Any)', fontsize = 10)
    plt.grid(True)
    plt.savefig('ec_2_2_Status_Check_Failed(Any).png', dpi=100)
    plt.close()
    
    #------------------PowerBI Gateway EC2 Instance Status Check Failed(Instance)----------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    ec_2_3_response = client.get_metric_statistics(
        Namespace='AWS/EC2',
        MetricName='StatusCheckFailed_Instance',
        Dimensions=[
            {
                'Name': 'InstanceId',
                'Value': 'i-0cc329b29763639a0'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(ec_2_3_response)
    ec_2_3_a = (ec_2_3_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    ec_2_3_date_list = []
    ec_2_3_cpu_list = []
    for i in ec_2_3_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        ec_2_3_date_list.append(est_dt)
        ec_2_3_cpu_list.append(c)
        
    ec_2_3_d = {"Dates : " : ec_2_3_date_list, "Status Check Failed(Instance) : " : ec_2_3_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    ec_2_3_df = pd.DataFrame(ec_2_3_d, columns=['Dates : ', 'Status Check Failed(Instance) : '])
    ec_2_3_df.sort_values(by=['Dates : ','Status Check Failed(Instance) : '], inplace = True)
    ec_2_3_Timestamp = ec_2_3_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    ec_2_3_Timestamp_list = []
    
    for j in ec_2_3_Timestamp:
        time = j.split()
        time = time[1]
        ec_2_3_Timestamp_list.append(time)
    ec_2_3_CPU_Utilization = ec_2_3_df['Status Check Failed(Instance) : '].tolist()
    
    plt.rcParams['font.size'] = '7'
    
    fig43 = plt.plot(ec_2_3_Timestamp_list, ec_2_3_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Status Check Failed(Instance) for PowerBI Gateway EC2 Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Status Check Failed(Instance)', fontsize = 10)
    plt.grid(True)
    plt.savefig('ec_2_3_Status_Check_Failed(Instance).png', dpi=100)
    plt.close()
    
    #------------------PowerBI Gateway EC2 Instance Status Check Failed(System)------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    ec_2_4_response = client.get_metric_statistics(
        Namespace='AWS/EC2',
        MetricName='StatusCheckFailed_System',
        Dimensions=[
            {
                'Name': 'InstanceId',
                'Value': 'i-0cc329b29763639a0'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(ec_2_4_response)
    ec_2_4_a = (ec_2_4_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    ec_2_4_date_list = []
    ec_2_4_cpu_list = []
    for i in ec_2_4_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        ec_2_4_date_list.append(est_dt)
        ec_2_4_cpu_list.append(c)
        
    ec_2_4_d = {"Dates : " : ec_2_4_date_list, "Status Check Failed(System) : " : ec_2_4_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    ec_2_4_df = pd.DataFrame(ec_2_4_d, columns=['Dates : ', 'Status Check Failed(System) : '])
    ec_2_4_df.sort_values(by=['Dates : ','Status Check Failed(System) : '], inplace = True)
    ec_2_4_Timestamp = ec_2_4_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    ec_2_4_Timestamp_list = []
    
    for j in ec_2_4_Timestamp:
        time = j.split()
        time = time[1]
        ec_2_4_Timestamp_list.append(time)
    ec_2_4_CPU_Utilization = ec_2_4_df['Status Check Failed(System) : '].tolist()
    
    plt.rcParams['font.size'] = '7'
    
    fig44 = plt.plot(ec_2_4_Timestamp_list, ec_2_4_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Status Check Failed(System) for PowerBI Gateway EC2 Instance', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Status Check Failed(System)', fontsize = 10)
    plt.grid(True)
    plt.savefig('ec_2_4_Status_Check_Failed(System).png', dpi=100)
    plt.close()
    
    #------------------- OpenSearch/Elastic Search Total Nodes ---------------------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    os_1_5_response = client.get_metric_statistics(
        Namespace='AWS/ES',
        MetricName='Nodes',
        Dimensions=[
            {   
                'Name': 'DomainName',
                'Value': 'prod-amc-es'
                
            },
            {
                'Name': 'ClientId',
                'Value': '719309552406'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    #print(os_1_5_response)
    os_1_5_a = (os_1_5_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    os_1_5_date_list = []
    os_1_5_cpu_list = []
    for i in os_1_5_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        os_1_5_date_list.append(est_dt)
        os_1_5_cpu_list.append(c)
        
    os_1_5_d = {"Dates : " : os_1_5_date_list, "Total Nodes : " : os_1_5_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    os_1_5_df = pd.DataFrame(os_1_5_d, columns=['Dates : ', 'Total Nodes : '])
    os_1_5_df.sort_values(by=['Dates : ','Total Nodes : '], inplace = True)
    os_1_5_Timestamp = os_1_5_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    os_1_5_Timestamp_list = []
    
    for j in os_1_5_Timestamp:
        time = j.split()
        time = time[1]
        os_1_5_Timestamp_list.append(time)
    os_1_5_CPU_Utilization = os_1_5_df['Total Nodes : '].tolist()
    
    fig45 = plt.rcParams['font.size'] = '7'
    
    plt.plot(os_1_5_Timestamp_list, os_1_5_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Total Nodes for OpenSearch/Elastic Search Cluster', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Total Nodes', fontsize = 10)
    plt.grid(True)
    plt.savefig('os_1_5_Total_Nodes.png', dpi=100)
    plt.close()
    
    
    #---------------------OpenSearch/Elastic Search Free Storage Space ---------------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    os_1_6_response = client.get_metric_statistics(
        Namespace='AWS/ES',
        MetricName="FreeStorageSpace",
        Dimensions=[
            {   
                'Name': 'DomainName',
                'Value': 'prod-amc-es'
                
            },
            {
                'Name': 'ClientId',
                'Value': '719309552406'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    os_1_6_a = (os_1_6_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    os_1_6_date_list = []
    os_1_6_cpu_list = []
    for i in os_1_6_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        os_1_6_date_list.append(est_dt)
        os_1_6_cpu_list.append(c)
        
    os_1_6_d = {"Dates : " : os_1_6_date_list, "Free Storage Space : " : os_1_6_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    os_1_6_df = pd.DataFrame(os_1_6_d, columns=['Dates : ', 'Free Storage Space : '])
    os_1_6_df.sort_values(by=['Dates : ','Free Storage Space : '], inplace = True)
    os_1_6_Timestamp = os_1_6_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    os_1_6_Timestamp_list = []
    
    for j in os_1_6_Timestamp:
        time = j.split()
        time = time[1]
        os_1_6_Timestamp_list.append(time)
    os_1_6_CPU_Utilization = os_1_6_df['Free Storage Space : '].tolist()
    
    plt.rcParams['font.size'] = '7'
    
    plt.plot(os_1_6_Timestamp_list, os_1_6_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Free Storage Space for OpenSearch/Elastic Search Cluster', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Free Storage Space', fontsize = 10)
    plt.grid(True)
    plt.savefig('os_1_6_Free_Storage_Space.png', dpi=100)
    plt.close()
    
    #----------------OpenSearch/Elastic Search KPI - Indexing Rate ------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    os_2_1_response = client.get_metric_statistics(
        Namespace='AWS/ES',
        MetricName='IndexingRate',
        Dimensions=[
            {   
                'Name': 'DomainName',
                'Value': 'prod-amc-es'
                
            },
            {
                'Name': 'ClientId',
                'Value': '719309552406'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    
    os_2_1_a = (os_2_1_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    os_2_1_date_list = []
    os_2_1_cpu_list = []
    for i in os_2_1_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        os_2_1_date_list.append(est_dt)
        os_2_1_cpu_list.append(c)
        
    os_2_1_d = {"Dates : " : os_2_1_date_list, "Indexing Rate : " : os_2_1_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    os_2_1_df = pd.DataFrame(os_2_1_d, columns=['Dates : ', 'Indexing Rate : '])
    os_2_1_df.sort_values(by=['Dates : ','Indexing Rate : '], inplace = True)
    os_2_1_Timestamp = os_2_1_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    os_2_1_Timestamp_list = []
    
    for j in os_2_1_Timestamp:
        time = j.split()
        time = time[1]
        os_2_1_Timestamp_list.append(time)
    os_2_1_CPU_Utilization = os_2_1_df['Indexing Rate : '].tolist()
    
    plt.rcParams['font.size'] = '7'
    
    plt.plot(os_2_1_Timestamp_list, os_2_1_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Indexing Rate for OpenSearch Key Performance Indicators', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Indexing Rate', fontsize = 10)
    plt.grid(True)
    plt.savefig('os_2_1_Indexing_Rate.png', dpi=100)
    plt.close()
    
    #------------------OpenSearch/Elastic Search KPI - Search Rate ------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    os_2_2_response = client.get_metric_statistics(
        Namespace='AWS/ES',
        MetricName='SearchRate',
        Dimensions=[
            {   
                'Name': 'DomainName',
                'Value': 'prod-amc-es'
                
            },
            {
                'Name': 'ClientId',
                'Value': '719309552406'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    
    os_2_2_a = (os_2_2_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    os_2_2_date_list = []
    os_2_2_cpu_list = []
    for i in os_2_2_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        os_2_2_date_list.append(est_dt)
        os_2_2_cpu_list.append(c)
        
    os_2_2_d = {"Dates : " : os_2_2_date_list, "Search Rate : " : os_2_2_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    os_2_2_df = pd.DataFrame(os_2_2_d, columns=['Dates : ', 'Search Rate : '])
    os_2_2_df.sort_values(by=['Dates : ','Search Rate : '], inplace = True)
    os_2_2_Timestamp = os_2_2_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    os_2_2_Timestamp_list = []
    
    for j in os_2_2_Timestamp:
        time = j.split()
        time = time[1]
        os_2_2_Timestamp_list.append(time)
    os_2_2_CPU_Utilization = os_2_2_df['Search Rate : '].tolist()
    
    plt.rcParams['font.size'] = '7'
    
    plt.plot(os_2_2_Timestamp_list, os_2_2_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Search Rate for OpenSearch Key Performance Indicators', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Search Rate', fontsize = 10)
    plt.grid(True)
    plt.savefig('os_2_2_Search_Rate.png', dpi=100)
    plt.close()
    
    #------------OpenSearch/Elastic Search KPI - Indexing Latency-------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    os_2_3_response = client.get_metric_statistics(
        Namespace='AWS/ES',
        MetricName='IndexingLatency',
        Dimensions=[
            {   
                'Name': 'DomainName',
                'Value': 'prod-amc-es'
                
            },
            {
                'Name': 'ClientId',
                'Value': '719309552406'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    
    os_2_3_a = (os_2_3_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    os_2_3_date_list = []
    os_2_3_cpu_list = []
    for i in os_2_3_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        os_2_3_date_list.append(est_dt)
        os_2_3_cpu_list.append(c)
        
    os_2_3_d = {"Dates : " : os_2_3_date_list, "Indexing Latency : " : os_2_3_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    os_2_3_df = pd.DataFrame(os_2_3_d, columns=['Dates : ', 'Indexing Latency : '])
    os_2_3_df.sort_values(by=['Dates : ','Indexing Latency : '], inplace = True)
    os_2_3_Timestamp = os_2_3_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    os_2_3_Timestamp_list = []
    
    for j in os_2_3_Timestamp:
        time = j.split()
        time = time[1]
        os_2_3_Timestamp_list.append(time)
    os_2_3_CPU_Utilization = os_2_3_df['Indexing Latency : '].tolist()
    
    plt.rcParams['font.size'] = '7'
    
    plt.plot(os_2_3_Timestamp_list, os_2_3_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Indexing Latency for OpenSearch Key Performance Indicators', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Indexing Latency', fontsize = 10)
    plt.grid(True)
    plt.savefig('os_2_3_Indexing_Latency.png', dpi=100)
    plt.close()
    
    #--------------------OpenSearch/Elastic Search KPI - Search Latency ------------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    os_2_4_response = client.get_metric_statistics(
        Namespace='AWS/ES',
        MetricName='SearchLatency',
        Dimensions=[
            {   
                'Name': 'DomainName',
                'Value': 'prod-amc-es'
                
            },
            {
                'Name': 'ClientId',
                'Value': '719309552406'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    
    os_2_4_a = (os_2_4_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    os_2_4_date_list = []
    os_2_4_cpu_list = []
    for i in os_2_4_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        os_2_4_date_list.append(est_dt)
        os_2_4_cpu_list.append(c)
        
    os_2_4_d = {"Dates : " : os_2_4_date_list, "Search Latency : " : os_2_4_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    os_2_4_df = pd.DataFrame(os_2_4_d, columns=['Dates : ', 'Search Latency : '])
    os_2_4_df.sort_values(by=['Dates : ','Search Latency : '], inplace = True)
    os_2_4_Timestamp = os_2_4_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    os_2_4_Timestamp_list = []
    
    for j in os_2_4_Timestamp:
        time = j.split()
        time = time[1]
        os_2_4_Timestamp_list.append(time)
    os_2_4_CPU_Utilization = os_2_4_df['Search Latency : '].tolist()
    
    plt.rcParams['font.size'] = '7'
    
    plt.plot(os_2_4_Timestamp_list, os_2_4_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Search Latency for OpenSearch Key Performance Indicators', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Search Latency', fontsize = 10)
    plt.grid(True)
    plt.savefig('os_2_4_Search_Latency.png', dpi=100)
    plt.close()
    
    #----------------OpenSearch/Elastic Search KPI - HTTP Requests by Response Code -----------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    os_2_5_response = client.get_metric_statistics(
        Namespace='AWS/ES',
        MetricName='4xx',
        Dimensions=[
            {   
                'Name': 'DomainName',
                'Value': 'prod-amc-es'
                
            },
            {
                'Name': 'ClientId',
                'Value': '719309552406'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    
    os_2_5_a = (os_2_5_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    os_2_5_date_list = []
    os_2_5_cpu_list = []
    for i in os_2_5_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        os_2_5_date_list.append(est_dt)
        os_2_5_cpu_list.append(c)
        
    os_2_5_d = {"Dates : " : os_2_5_date_list, "HTTP Requests by Response Code : " : os_2_5_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    os_2_5_df = pd.DataFrame(os_2_5_d, columns=['Dates : ', 'HTTP Requests by Response Code : '])
    os_2_5_df.sort_values(by=['Dates : ','HTTP Requests by Response Code : '], inplace = True)
    os_2_5_Timestamp = os_2_5_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    os_2_5_Timestamp_list = []
    
    for j in os_2_5_Timestamp:
        time = j.split()
        time = time[1]
        os_2_5_Timestamp_list.append(time)
    os_2_5_CPU_Utilization = os_2_5_df['HTTP Requests by Response Code : '].tolist()
    
    plt.rcParams['font.size'] = '7'
    
    plt.plot(os_2_5_Timestamp_list, os_2_5_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('HTTP Requests by Response Code for OpenSearch Key Performance Indicators', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('HTTP Requests by Response Code', fontsize = 10)
    plt.grid(True)
    plt.savefig('os_2_5_HTTP_Requests_by_Response_Code.png', dpi=100)
    plt.close()
    
    #---------------------- OpenSearch/Elastic Search KPI - Invalid Host Header Request --------------------
    
    client = boto3.client('cloudwatch',aws_access_key_id=ac_id,aws_secret_access_key=as_key, aws_session_token = a_s_t, region_name='us-west-2')
    
    os_2_6_response = client.get_metric_statistics(
        Namespace='AWS/ES',
        MetricName='ElasticsearchRequests',
        Dimensions=[
            {   
                'Name': 'DomainName',
                'Value': 'prod-amc-es'
                
            },
            {
                'Name': 'ClientId',
                'Value': '719309552406'
            }
        ],
        StartTime=datetime(y_input_year,y_input_month,y_input_day,22,0),
        EndTime=datetime(t_input_year,t_input_month,t_input_day,10,30),
        Period=3600,
        Statistics=['Average']
    )
    
    os_2_6_a = (os_2_6_response['Datapoints'])
    
    #Code to convert UTC Timezone to EST Timezone. 
    
    os_2_6_date_list = []
    os_2_6_cpu_list = []
    for i in os_2_6_a:
        b = i['Timestamp']
        c = i['Average']
        utc_dt = b
        utc_dt.strftime(fmt)
        est_dt = utc_dt.astimezone(est).strftime(fmt)
        os_2_6_date_list.append(est_dt)
        os_2_6_cpu_list.append(c)
        
    os_2_6_d = {"Dates : " : os_2_6_date_list, "Invalid Host Header Requests : " : os_2_6_cpu_list}
    
    #Code to Sort the Statistics values based on Time using Pandas.
    
    os_2_6_df = pd.DataFrame(os_2_6_d, columns=['Dates : ', 'Invalid Host Header Requests : '])
    os_2_6_df.sort_values(by=['Dates : ','Invalid Host Header Requests : '], inplace = True)
    os_2_6_Timestamp = os_2_6_df['Dates : '].tolist()
    
    #Code to convert Dataframe to List and then to String
    
    os_2_6_Timestamp_list = []
    
    for j in os_2_6_Timestamp:
        time = j.split()
        time = time[1]
        os_2_6_Timestamp_list.append(time)
    os_2_6_CPU_Utilization = os_2_6_df['Invalid Host Header Requests : '].tolist()
    
    plt.rcParams['font.size'] = '7'
    
    plt.plot(os_2_6_Timestamp_list, os_2_6_CPU_Utilization, color = 'blue', marker = 'o')
    
    plt.title('Invalid Host Header Requests for OpenSearch Key Performance Indicators', fontsize = 10)
    plt.xlabel('Time(in EST for last 12 Hours)', fontsize = 10)
    plt.ylabel('Invalid Host Header Requests', fontsize = 10)
    plt.grid(True)
    plt.savefig('os_2_6_Invalid_Host_Header_Requests.png', dpi=100)
    plt.close()
    
    
    #-------------Resources CPU Utilizations ---------------
    
    RDS_Writer_CPU = rds_1_1_df.tail(1)
    RDS_Reader_CPU = rds_2_1_df.tail(1)
    Redshift_Writer_Instance_CPU = rs_1_1_df.tail(1)
    Redshift_Reader_Instance_CPU = rs_2_1_df.tail(1)
    Document_DB_Writer_CPU = db_1_1_df.tail(1)
    Document_DB_Reader_CPU = db_2_1_df.tail(1)
    SGW_EC2_CPU = ec_1_1_df.tail(1)
    PBI_EC2_CPU = ec_2_1_df.tail(1)
    
    
    
    rds1 = (RDS_Writer_CPU.to_string())
    rds2 = (RDS_Reader_CPU.to_string())
    rs1 = (Redshift_Writer_Instance_CPU.to_string())
    rs2 = (Redshift_Reader_Instance_CPU.to_string())
    db1 = (Document_DB_Writer_CPU.to_string())
    db2 = (Document_DB_Reader_CPU.to_string())
    ec1 = (SGW_EC2_CPU.to_string())
    ec2 = (PBI_EC2_CPU.to_string())
    
    rds1_list = rds1.split()
    rds2_list = rds2.split()
    rs1_list = rs1.split()
    rs2_list = rs2.split()
    db1_list = db1.split()
    db2_list = db2.split()
    ec1_list = ec1.split()
    ec2_list = ec2.split()
    
    rds1_cpu = rds1_list[-1]
    rds2_cpu = rds2_list[-1]
    rs1_cpu = rs1_list[-1]
    rs2_cpu = rs2_list[-1]
    db1_cpu = db1_list[-1]
    db2_cpu = db2_list[-1]
    ec1_cpu = ec1_list[-1]
    ec2_cpu = ec2_list[-1]
    
    
    
    #------------------ Adding Images to Document ------------- 
    
    """document = Document()
    
    p = document.add_paragraph('Daily Status Monitoring Report'+str(t_input_month)+"/"+str(t_input_day)+"/"+str(t_input_year))
    r = p.add_run()
    r.alignment = 1
    r.add_picture('db_1_1_CPU_Utilization.png')
    
    
    document.save('amc_mon_report.docx')"""
    
    document = Document()
    
    document.add_heading('Daily Status Monitoring Report - ' +str(t_input_month)+"/"+str(t_input_day)+"/"+str(t_input_year) , 1)
    
    # --------------CPU Utilization Table Details ------------------
    
    
    """cpu_values = [[1, "Aurora Postgres Writer Instance", rds1_cpu], [2, "Aurora Postgres Reader Instance", rds2_cpu],
    [3, "Redshift Cluster(americold-redshift-prod)", rs1_cpu], [4, "Redshift Cluster(americold-redshift-prod-readonly)", rs2_cpu],
    [5, "Document DB Primary Instance", db1_cpu], [6, "Document DB Replica Instance", db2_cpu],
    [7, "EC2 Storage Gateway Instance", ec1_cpu], [8, "EC2 PowerBI Gateway Instance", ec2_cpu]]
    
    cpu_table = document.add_table(rows=1, cols=4)
    cpu_table.style = "Table Grid"
    hdr_cells = cpu_table.rows[0].cells
    hdr_cells[0].text = "S_no"
    hdr_cells[1].text = "Name_of_the_Resource"
    hdr_cells[2].text = "CPU_Utilization"
    hdr_cells[3].text = "Status"
    
    for S_no, Name_of_the_Resource, CPU_Utilization in cpu_values:
        row_cells = cpu_table.add_row().cells
        row_cells[0].text = str(S_no)
        row_cells[1].text = Name_of_the_Resource
        row_cells[2].text = CPU_Utilization
        
    """   
    
    
    p = document.add_paragraph()
    r = p.add_run()
    r.alignment = 1
    r.bold = True
    
    
    #----RDS Metrics ------------
    #r.add_text('\n\n')
    r.add_text('CPU Utilization of RDS Aurora PostgresSQL(Writer)')
    r.add_picture(r'rds_1_1_CPU.png')
    r.add_text('DB Connections for RDS Aurora PostgresSQL(Writer)')
    r.add_picture(r'rds_1_2_DB_Connection.png')
    r.add_text('Freeable Memory for RDS Aurora PostgresSQL(Writer)')
    r.add_picture(r'rds_1_3_Freeable_Memory.png')
    r.add_text('Write IOPS for RDS Aurora PostgresSQL(Writer)')
    r.add_picture(r'rds_1_4_Write_IOPS.png')
    r.add_text('Reader IOPS for RDS Aurora PostgresSQL(Writer)')
    r.add_picture(r'rds_1_5_Reader_IOPS.png')
    r.add_text('Queue Depth for RDS Aurora PostgresSQL(Writer)')
    r.add_picture(r'rds_1_6_Queue_Depth.png')
    r.add_text('CPU Utilization of RDS Aurora PostgresSQL(Reader)')
    r.add_picture(r'rds_2_1_CPU_Utilization.png')
    r.add_text('DB Connections for RDS Aurora PostgresSQL(Reader)')
    r.add_picture(r'rds_2_2_DB_Connections.png')
    r.add_text('Freeable Memory for RDS Aurora PostgresSQL(Reader)')
    r.add_picture(r'rds_2_3_Freeable_Memory.png')
    r.add_text('Write IOPS for RDS Aurora PostgresSQL(Reader)')
    r.add_picture(r'rds_2_4_Write_IOPS.png')
    r.add_text('Reader IOPS for RDS Aurora PostgresSQL(Reader)')
    r.add_picture(r'rds_2_5_Reader_IOPS.png')
    r.add_text('Queue Depth for RDS Aurora PostgresSQL(Reader)')
    r.add_picture(r'rds_2_6_Queue_Depth.png')
    
    #---------Redshift Metrics -----------
    
    r.add_text('CPU Utilization for Reshift Writer Instance')
    r.add_picture(r'rs_1_1_CPU_Utilization.png')
    r.add_text('% of Diskspace used for Reshift Writer Instance')
    r.add_picture(r'rs_1_2_%diskspace_used.png')
    r.add_text('Health Status for Reshift Writer Instance')
    r.add_picture(r'rs_1_3_Health_Status.png')
    r.add_text('CPU Utilization for Reshift Reader Instance')
    r.add_picture(r'rs_2_1_CPU_Utilization.png')
    r.add_text('% of Diskspace used for Reshift Reader Instance')
    r.add_picture(r'rs_2_2_%diskspace_used.png')
    r.add_text('Health Status for Reshift Reader Instance')
    r.add_picture(r'rs_2_3_Health_Status.png')
    
    #----------- Storage Gateway Metrics ------------
    
    r.add_text('Cache Percent Dirty for Storage Gateway Instance')
    r.add_picture(r'sg_3_Cache_Percent_Dirty.png')
    r.add_text('Cache Hit Dirty for Storage Gateway Instance')
    r.add_picture(r'sg_4_Cache_Hit_Dirty.png')
    
    #------------- Document DB Metrics --------------
    
    r.add_text('CPU Utilization for Document DB Writer Instance')
    r.add_picture(r'db_1_1_CPU_Utilization.png')
    r.add_text('Freeable Memory for Document DB Writer Instance')
    r.add_picture(r'db_1_2_Freeable_Memory.png')
    r.add_text('Free Local Storage for Document DB Writer Instance')
    r.add_picture(r'db_1_3_Free_Local_Storage.png')
    r.add_text('Swap Usage for Document DB Writer Instance')
    r.add_picture(r'db_1_4_Swap_Usage.png')
    r.add_text('Database Connections for Document DB Writer Instance')
    r.add_picture(r'db_1_5_Database_Connections.png')
    r.add_text('Maximum Database Connections for Document DB Writer Instance')
    r.add_picture(r'db_1_6_Maximum_Database_Connections.png')
    r.add_text('Database Cursors for Document DB Writer Instance')
    r.add_picture(r'db_1_7_Database_Cursors.png')
    r.add_text('Maximum Database Cursors for Document DB Writer Instance')
    r.add_picture(r'db_1_8_Maximum_Database_Cursors.png')
    
    r.add_text('CPU Utilization for Document DB Reader Instance')
    r.add_picture(r'db_2_1_CPU_Utilization.png')
    r.add_text('Freeable Memory for Document DB Reader Instance')
    r.add_picture(r'db_2_2_Freeable_Memory.png')
    r.add_text('Free Local Storage for Document DB Reader Instance')
    r.add_picture(r'db_2_3_Free_Local_Storage.png')
    r.add_text('Swap Usage for Document DB Reader Instance')
    r.add_picture(r'db_2_4_Swap_Usage.png')
    r.add_text('Database Connections for Document DB Reader Instance')
    r.add_picture(r'db_2_5_Database_Connections.png')
    r.add_text('Maximum Database Connections for Document DB Reader Instance')
    r.add_picture(r'db_2_6_Maximum_Database_Connections.png')
    r.add_text('Database Cursors for Document DB Reader Instance')
    r.add_picture(r'db_2_7_Database_Cursors.png')
    r.add_text('Maximum Database Cursors for Document DB Reader Instance')
    r.add_picture(r'db_2_8_Maximum_Database_Cursors.png')
    
    #--------------- EC2 Metrics ------------------
    
    r.add_text('CPU Utilization for Storage Gateway EC2 Instance')
    r.add_picture(r'ec_1_1_CPU_Utilization.png')
    r.add_text('Status Check Failed(Any) for Storage Gateway EC2 Instance')
    r.add_picture(r'ec_1_2_Status_Check_Failed(Any).png')
    r.add_text('Status Check Failed(Instance) for Storage Gateway EC2 Instance')
    r.add_picture(r'ec_1_3_Status_Check_Failed(Instance).png')
    r.add_text('Status Check Failed(System) for Storage Gateway EC2 Instance')
    r.add_picture(r'ec_1_4_Status_Check_Failed(System).png')
    r.add_text('CPU Utilization for PowerBI Gateway EC2 Instance')
    r.add_picture(r'ec_2_1_CPU_Utilization.png')
    r.add_text('Status Check Failed(Any) for PowerBI Gateway EC2 Instance')
    r.add_picture(r'ec_2_2_Status_Check_Failed(Any).png')
    r.add_text('Status Check Failed(Instance) for PowerBI Gateway EC2 Instance')
    r.add_picture(r'ec_2_3_Status_Check_Failed(Instance).png')
    r.add_text('Status Check Failed(System) for PowerBI Gateway EC2 Instance')
    r.add_picture(r'ec_2_4_Status_Check_Failed(System).png')
    
    #------------- OpenSearch Metrics ----------------
    
    r.add_text('Total Nodes for OpenSearch/Elastic Search Cluster')
    r.add_picture(r'os_1_5_Total_Nodes.png')
    r.add_text('Free Storage Space for OpenSearch/Elastic Search Cluster')
    r.add_picture(r'os_1_6_Free_Storage_Space.png')
    r.add_text('Indexing Rate for OpenSearch Key Performance Indicators')
    r.add_picture(r'os_2_1_Indexing_Rate.png')
    r.add_text('Search Rate for OpenSearch Key Performance Indicators')
    r.add_picture(r'os_2_2_Search_Rate.png')
    r.add_text('Indexing Latency for OpenSearch Key Performance Indicators')
    r.add_picture(r'os_2_3_Indexing_Latency.png')
    r.add_text('Search Latency for OpenSearch Key Performance Indicators')
    r.add_picture(r'os_2_4_Search_Latency.png')
    r.add_text('HTTP Requests by Response Code for OpenSearch Key Performance Indicators')
    r.add_picture(r'os_2_5_HTTP_Requests_by_Response_Code.png')
    r.add_text('Invalid Host Header Requests for OpenSearch Key Performance Indicators')
    r.add_picture(r'os_2_6_Invalid_Host_Header_Requests.png')
    
    
    r.add_text(' END OF THE DOC ')
    
    document.save('DSMR.docx')
    
    
    
    #----------------- Sending Email ----------------------------
    
    def gmail_login(user_name, pass_word):
        smtp_server_name = "smtp.gmail.com"
        port_number = 587 
        server = smtplib.SMTP(host=smtp_server_name, port=port_number)
        server.connect(host=smtp_server_name, port=port_number)
        server.starttls()
        server.login(user=user_name, password=pass_word)
        return server
    
    msg = MIMEMultipart()
    msg['Subject'] = "Daily Status Monitoring Report - "+str(t_input_month)+"/"+str(t_input_day)+"/"+str(t_input_year) +" | PROD"
    msg['From'] = 'amcstatusmail@gmail.com'
    #msg['To'] = 'ddoddi@miraclesoft.com'
    msg['To'] = 'dilipkumar.doddi@americold.com'
    
    # - - - - - CPU Values - - - - -
    
    rds1_cpu_val = round(float(rds1_cpu),3)
    rds2_cpu_val = round(float(rds2_cpu),3)
    rs1_cpu_val = round(float(rs1_cpu),3)
    rs2_cpu_val = round(float(rs2_cpu),3)
    db1_cpu_val = round(float(db1_cpu),3)
    db2_cpu_val = round(float(db2_cpu),3)
    ec1_cpu_val = round(float(ec1_cpu),3)
    ec2_cpu_val = round(float(ec2_cpu),3)
    
    # - - - - - CPU Status - - - - -
    
    rds1_cpu_stat = ""
    rds2_cpu_stat = ""
    rs1_cpu_stat = ""
    rs2_cpu_stat = ""
    db1_cpu_stat = ""
    db2_cpu_stat = ""
    ec1_cpu_stat = ""
    ec2_cpu_stat = ""
    
    
    
    if rds1_cpu_val >= 75.000:
        rds1_cpu_stat = '<h4 style="color:red;"><b>Unhealthy</b></h4>'
    else:
        rds1_cpu_stat = '<h4 style="color:green;">Healthy</h4>'
    
    if rds2_cpu_val >= 75.000:
        rds2_cpu_stat = '<h4 style="color:red;"><b>Unhealthy</b></h4>'
    else:
        rds2_cpu_stat = '<h4 style="color:green;">Healthy</h4>'
        
    if rs1_cpu_val >= 75.000:
        rs1_cpu_stat = '<h4 style="color:red;"><b>Unhealthy</b></h4>'
    else:
        rs1_cpu_stat = '<h4 style="color:green;">Healthy</h4>'
        
    if rs2_cpu_val >= 75.000:
        rs2_cpu_stat = '<h4 style="color:red;"><b>Unhealthy</b></h4>'
    else:
        rs2_cpu_stat = '<h4 style="color:green;">Healthy</h4>'
        
    if db1_cpu_val >= 75.000:
        db1_cpu_stat = '<h4 style="color:red;"><b>Unhealthy</b></h4>'
    else:
        db1_cpu_stat = '<h4 style="color:green;">Healthy</h4>'
        
    if db2_cpu_val >= 75.000:
        db2_cpu_stat = '<h4 style="color:red;"><b>Unhealthy</b></h4>'
    else:
        db2_cpu_stat = '<h4 style="color:green;">Healthy</h4>'
        
    if ec1_cpu_val >= 75.000:
        ec1_cpu_stat = '<h4 style="color:red;"><b>Unhealthy</b></h4>'
    else:
        ec1_cpu_stat = '<h4 style="color:green;">Healthy</h4>'
    
    if ec2_cpu_val >= 75.000:
        ec2_cpu_stat = '<h4 style="color:red;"><b>Unhealthy</b></h4>'
    else:
        ec2_cpu_stat = '<h4 style="color:green;">Healthy</h4>'
    
    
    
    body = """
    <h4 style="color:red"><b>NOTE:</b></h4> <p><b>Please Do Not Reply to this Email. If you have any queries please reply back to DilipKumar.Doddi@americold.com </b></p>
    <br><br>
    Hi Team, <br><br> Please find the attachment.<br><br> 
    
    <h2>AWS Resources CPU Utilization</h2>
    
    <table>
      <tr>
        <th>S. no</th>
        <th>Name of the Resource</th>
        <th>CPU Utilization</th>
        <th>Health Status</th>
      </tr>
      <tr>
        <td>1. </td>
        <td>Aurora Postgres Writer Instance</td>
        <td><h4>{rds1_cpu_val}</h4></td>
        <td>
        <h4>{rds1_cpu_stat}</h4>
        </td>
      </tr>
      <tr>
        <td>2. </td>
        <td>Aurora Postgres Reader Instance</td>
        <td><h4>{rds2_cpu_val}</h4></td>
        <td>
        <h4>{rds2_cpu_stat}</h4>
        </td>
      </tr>
      <tr>
        <td>3. </td>
        <td>Redshift Cluster(americold-redshift-prod)</td>
        <td><h4>{rs1_cpu_val}</h4></td>
        <td>
        <h4>{rs1_cpu_stat}</h4>
        </td>
      </tr>
      <tr>
        <td>4. </td>
        <td>Redshift Cluster(americold-redshift-prod-readonly)</td>
        <td><h4>{rs2_cpu_val}</h4></td>
        <td>
        <h4>{rs2_cpu_stat}</h4>
        </td>
      </tr>
      <tr>
        <td>5. </td>
        <td>Document DB Primary Instance</td>
        <td><h4>{db1_cpu_val}</h4></td>
        <td>
        <h4>{db1_cpu_stat}</h4>
        </td>
      </tr>
      <tr>
        <td>6. </td>
        <td>Document DB Replica Instance</td>
        <td><h4>{db2_cpu_val}</h4></td>
        <td>
        <h4>{db2_cpu_stat}</h4>
        </td>
      </tr>
      <tr>
        <td>7. </td>
        <td>EC2 Storage Gateway Instance</td>
        <td><h4>{ec1_cpu_val}</h4></td>
        <td>
        <h4>{ec1_cpu_stat}</h4>
        </td>
      </tr>
      <tr>
        <td>8. </td>
        <td>EC2 PowerBI Gateway Instance</td>
        <td><h4>{ec2_cpu_val}</h4></td>
        <td>
        <h4>{ec2_cpu_stat}</h4> 
        </td>
      </tr>
    </table>
    
     <br><br>Regards,<br> Dilip Kumar
    """.format(rds1_cpu_val=rds1_cpu_val, rds2_cpu_val=rds2_cpu_val, rs1_cpu_val=rs1_cpu_val, rs2_cpu_val=rs2_cpu_val, db1_cpu_val=db1_cpu_val, db2_cpu_val=db2_cpu_val, ec1_cpu_val=ec1_cpu_val, ec2_cpu_val=ec2_cpu_val, rds1_cpu_stat = rds1_cpu_stat, rds2_cpu_stat = rds2_cpu_stat, rs1_cpu_stat = rs1_cpu_stat, rs2_cpu_stat = rs2_cpu_stat, db1_cpu_stat = db1_cpu_stat, db2_cpu_stat = db2_cpu_stat, ec1_cpu_stat = ec1_cpu_stat, ec2_cpu_stat = ec2_cpu_stat)
    
    msg.attach(MIMEText(body, 'html'))
    
    file_path = r"C:\Users\ddoddi\Desktop\Notepad++\New5"
    file_name = 'DSMR.docx'
    
    file = open(file_path+"\\"+file_name, "rb")
    
    payload = MIMEBase("application", "octet-stream")
    payload.set_payload(file.read())
    file.close()
    encoders.encode_base64(payload)
    payload.add_header('Content-Disposition', 'attachment', filename=file_name)
    msg.attach(payload)
    
    server = gmail_login(user_name = 'amcstatusmail@gmail.com', pass_word = 'fcdxrlcxmfpadlqh')
    server.send_message(msg)
    server.quit()
    
    
    print("Email has been successfully sent!")
    
    return {
        'statusCode': 200,
        'body': json.dumps('Hello World')
    }
    